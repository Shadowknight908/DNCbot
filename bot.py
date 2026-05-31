"""DNC Lore Bot v0.8 — main entry point.

Major additions vs v0.5.1:
- Reply support for chat, chatuc, queries, and GM rulings (5-msg context cap)
- !DNC chatuc — admin-only "unhinged" mode with separate prompt
- !DNC gm <message-link> — GM ruling adjudication mode
- Role-based permissions (admin_roles, gm_roles in config)
- Hybrid retrieval: queries see summary + full original text
- GM revisions: replies from gm_roles trigger ruling revision
"""
from __future__ import annotations

import asyncio
import io
import json
import logging
import os
import random
import re
import sys
from datetime import datetime, timezone
from typing import List, Optional, Tuple

import discord
import yaml
from discord.ext import commands, tasks
from dotenv import load_dotenv

from channel_names import (
    normalize_channel_name,
    sanitize_channel_list,
    emit_warnings,
)
from espionage_store import EspionageStore
from war_store import WarStore
from chat_blacklist import ChatBlacklist
from inference_client import InferenceClient
from file_logging import FileLoggers
from memory_store import MemoryStore
from optout_store import OptOutStore
from prompt_store import PromptStore
from state_store import StateStore
from tavily_client import TavilyClient
import reply_chain
import year_scheduler

from map_colors import MapColorStore
from faction_store import FactionStore
from map_store import MapStore
import map_renderer
import map_scheduler as map_sched
from nickname_parser import parse_tag
from province_store import ProvinceStore
from map_cache import MapCache
import map_llm

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)
log = logging.getLogger("dnc")


# ----------------------------------------------------------------------
# ANSI colors for console output
_C_HDR  = "\033[1;37m"   # bold white  — section dividers + INGEST header
_C_SEND = "\033[96m"     # bright cyan — outbound → requests
_C_RECV = "\033[36m"     # cyan        — inbound  ← responses
_C_OK   = "\033[92m"     # bright green — success / store
_C_ERR  = "\033[91m"     # bright red   — errors
_C_DIM  = "\033[2m"      # dim          — supplementary / skipped
_C_RST  = "\033[0m"


def _fmt_think(v) -> str:
    """Human-readable label for a thinking_budget config value."""
    if v is None:
        return "default"
    if isinstance(v, str):
        return f"effort:{v}"
    if v == 0:
        return "off"
    return f"{v}tok"


# ----------------------------------------------------------------------
MENTION_RE = re.compile(r"<@!?(\d+)>")
MSG_LINK_RE = re.compile(
    r"https?://(?:\w+\.)?discord\.com/channels/(\d+)/(\d+)/(\d+)"
)
ID_RE = re.compile(r"^\d{15,21}$")


def load_config(path: str = "config.yaml") -> dict:
    with open(path, "r", encoding="utf-8") as f:
        return yaml.safe_load(f)


def parse_void_target(arg: str):
    arg = arg.strip()
    m = MSG_LINK_RE.match(arg)
    if m:
        return "message_link", {"guild_id": m.group(1), "channel_id": m.group(2),
                                "message_id": m.group(3)}
    m = MENTION_RE.match(arg)
    if m:
        return "mention", {"user_id": m.group(1)}
    if ID_RE.match(arg):
        return "snowflake", {"id": arg}
    return "unknown", {}


_DOC_CONTENT_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
}
_DOC_CHUNK_CHARS = 4500
_DOC_CHUNK_OVERLAP = 400
_ARCHIVIST_INPUT_CHARS = 60000
_EMBED_INPUT_CHARS = 30000
_GM_PRIOR_CONTEXT_CHARS = 150000
_GM_WIDER_CONTEXT_CHARS = 120000
_GM_CONTEXT_ITEM_CHARS = 30000

_ARCHIVIST_VALID_ACTION_TYPES = {
    "internal", "military", "economic", "diplomatic", "cultural", "other",
}


def _parse_archivist_output(raw: str) -> tuple[str, dict]:
    """Split the Archivist response into (summary, tags_dict).

    Expected shape:
        SUMMARY:
        <prose>

        TAGS:
        entities: ...
        action_type: ...
        claims: ...

    On any parse failure, returns (raw, {}). NO_LORE responses pass through
    as summary with empty tags.
    """
    tags: dict[str, str] = {"entities": "", "action_type": "", "claims": ""}
    if not raw:
        return "", tags

    if raw.strip().upper().startswith("NO_LORE"):
        return raw, tags

    # Locate TAGS: as its own line (case-sensitive, line-anchored)
    m = re.search(r"(?im)^\s*TAGS:\s*$", raw)
    if not m:
        # No tags block — strip a leading SUMMARY: line if present, return the rest
        stripped = re.sub(r"(?im)^\s*SUMMARY:\s*\n?", "", raw, count=1).strip()
        return stripped or raw.strip(), tags

    summary_section = raw[:m.start()]
    tags_section = raw[m.end():]
    summary = re.sub(r"(?im)^\s*SUMMARY:\s*\n?", "", summary_section, count=1).strip()
    if not summary:
        summary = raw.strip()

    for line in tags_section.splitlines():
        line = line.strip()
        if not line or ":" not in line:
            continue
        key, _, val = line.partition(":")
        key = key.strip().lower()
        val = val.strip()
        if key == "entities":
            tags["entities"] = val
        elif key == "action_type":
            v = val.lower().split()[0] if val else ""
            tags["action_type"] = v if v in _ARCHIVIST_VALID_ACTION_TYPES else ""
        elif key == "claims":
            tags["claims"] = val

    return summary, tags


def _doc_attachment_type(a: discord.Attachment) -> str | None:
    """Return 'pdf' or 'docx' if the attachment looks like a document, else None."""
    ct = (a.content_type or "").lower().split(";")[0].strip()
    if ct in _DOC_CONTENT_TYPES:
        return _DOC_CONTENT_TYPES[ct]
    fn = (a.filename or "").lower()
    if fn.endswith(".pdf"):
        return "pdf"
    if fn.endswith(".docx"):
        return "docx"
    return None


def _collapse_blank_lines(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+\n", "\n", text)
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def _trim_context_text(text: str, limit: int) -> str:
    if limit <= 0:
        return ""
    if len(text) <= limit:
        return text
    return text[:limit].rstrip() + f"\n\n[... truncated after {limit} chars ...]"


def _iter_text_chunks(text: str, chunk_chars: int = _DOC_CHUNK_CHARS,
                      overlap: int = _DOC_CHUNK_OVERLAP):
    clean = _collapse_blank_lines(text)
    if not clean:
        return
    start = 0
    while start < len(clean):
        end = min(len(clean), start + chunk_chars)
        yield clean[start:end].strip()
        if end >= len(clean):
            break
        start = max(0, end - overlap)


# ----------------------------------------------------------------------
class LoreBot(commands.Bot):
    def __init__(self, config: dict):
        intents = discord.Intents.default()
        intents.message_content = True
        intents.guilds = True
        intents.members = True
        
        self.prefix: str = config["bot"]["command_prefix"]
        super().__init__(command_prefix=self.prefix, intents=intents, help_command=None, strip_after_prefix=True)

        self.cfg = config

        # Sanitize channel name lists
        d_chan = config["discord"].get("channels", {})
        scan_clean, sw = sanitize_channel_list(d_chan.get("scan", []), "scan_channels")
        ignored_clean, iw = sanitize_channel_list(d_chan.get("ignored", []), "ignored_channels")
        admin_clean, aw = sanitize_channel_list(d_chan.get("admin", []), "admin_channels")
        
        gm_cfg = config.get("gm", {})
        gm_chan_raw, gm_cw = sanitize_channel_list(gm_cfg.get("channels", []), "gm_channels")
        emit_warnings(sw + iw + aw + gm_cw)

        self._scan_set = {normalize_channel_name(c) for c in scan_clean}
        self._ignored_set = {normalize_channel_name(c) for c in ignored_clean}
        self._admin_set = {normalize_channel_name(c) for c in admin_clean}
        self._gm_channel_set = {normalize_channel_name(c) for c in gm_chan_raw}
        self._gm_output_channel_name = normalize_channel_name(gm_cfg.get("output_channel", "") or "")
        self._scan_raw = scan_clean
        self._ignored_raw = ignored_clean
        self._admin_raw = admin_clean

        # Role configuration (case-insensitive matching)
        d_roles = config["discord"].get("roles", {})
        self._admin_roles = {
            r.strip().lower() for r in d_roles.get("admin", []) or []
            if r and isinstance(r, str)
        }
        self._gm_roles = {
            r.strip().lower() for r in gm_cfg.get("roles", []) or []
            if r and isinstance(r, str)
        }
        self._chatuc_roles = {
            r.strip().lower() for r in d_roles.get("chatuc", []) or []
            if r and isinstance(r, str)
        }

        # War GM mode configuration
        war_cfg = config.get("war", {})
        self._war_cfg = war_cfg
        self._war_roles = {
            r.strip().lower() for r in war_cfg.get("roles", []) or []
            if r and isinstance(r, str)
        }

        # Espionage channel sets
        esp_cfg = config.get("espionage", {})
        esp_cmd_raw, _ = sanitize_channel_list(
            esp_cfg.get("command_channels", []), "espionage_command_channels"
        )
        esp_sec_raw, _ = sanitize_channel_list(
            esp_cfg.get("secret_channels", []), "espionage_secret_channels"
        )
        self._espionage_cfg = esp_cfg
        self._espionage_command_set = {normalize_channel_name(c) for c in esp_cmd_raw}
        self._espionage_secret_set = {normalize_channel_name(c) for c in esp_sec_raw}

        # Conversation limits
        conv = config["bot"].get("conversation", {})
        self._max_chain_depth = int(conv.get("max_chain_depth", 5))
        self._max_history_chars = int(conv.get("max_history_message_chars", 2000))

        # Post-chain grouping: buffer consecutive same-author messages before ingesting
        self._chain_delay: float = float(config["bot"].get("chain_delay_seconds", 30))
        self._pending_chains: dict[tuple[int, int], list[discord.Message]] = {}
        self._pending_timers: dict[tuple[int, int], asyncio.Task] = {}

        # Maps ruling Discord message IDs → routing info for player rerolls via 🎲 reaction
        self._ruling_msg_cache: Dict[str, dict] = {}
        self._gm_inflight: set[str] = set()

        # War GM mode: maps war-ruling Discord message IDs → re-adjudication context
        # (used for 🎲 rerolls and GM revise-by-reply). In-memory, like the GM cache.
        self._war_ruling_cache: Dict[str, dict] = {}
        self._war_inflight: set[str] = set()

        # Subsystems
        m_cfg = config["models"]
        prov = m_cfg["provider"]
        main_api_key = os.environ["OPENROUTER_API_KEY"]
        embed_api_key = os.environ.get("EMBEDDING_API_KEY") or main_api_key
        self.llm = InferenceClient(
            api_key=main_api_key,
            base_url=prov["base_url"],
            site_url=prov.get("site_url", ""),
            site_name=prov.get("site_name", "DNC Lore Bot"),
            chat_model=m_cfg.get("defaults", {}).get("model", ""),
            embedding_model=m_cfg["embedding_model"],
            embedding_base_url=prov.get("embedding_base_url", ""),
            embedding_api_key=embed_api_key,
            vision_model=m_cfg.get("vision_model", ""),
        )
        self.memory = MemoryStore(config["memory"]["db_path"])
        self.state = StateStore("state.json")
        self.optouts = OptOutStore("optouts.txt")
        self.espionage = EspionageStore("espionage.json")
        self.war_store = WarStore("wars.json")
        self.chat_blacklist = ChatBlacklist(
            config.get("chat_blacklist_file", "chat_blacklist.txt")
        )
        self.flog = FileLoggers("logs")
        self.prompts = PromptStore(config.get("prompts", {}))

        map_cfg = config.get("spatial_mapping", {})
        self.map_colors = MapColorStore(
            path="map_colors.txt",
            palette=map_cfg.get("color_palette") or None,
        )
        self.map_store = MapStore(self.map_colors)
        self.faction_store = FactionStore(
            path=map_cfg.get("factions_path", "factions"),
            palette=map_cfg.get("color_palette") or None,
        )
        self.province_store = ProvinceStore(
            color_store=self.map_colors,
            definitions_path=map_cfg.get("province_data_path", "map_data/provinces.json"),
            ownership_path=map_cfg.get("province_ownership_path", "province_ownership.json"),
        )

        persist_path = (
            "map_data/last_world_map"
            if map_cfg.get("cache_persist", True)
            else None
        )
        self.map_cache = MapCache(persist_path=persist_path)

        tavily_key = os.environ.get("TAVILY_API_KEY", "")
        tavily_cfg = config.get("tavily", {})
        if tavily_key:
            self.tavily: Optional[TavilyClient] = TavilyClient(
                api_key=tavily_key,
                max_results=int(tavily_cfg.get("max_results", 5)),
                search_depth=tavily_cfg.get("search_depth", "basic"),
            )
        else:
            self.tavily = None
            enabled = tavily_cfg.get("enabled_modes", {})
            if any(enabled.values()):
                log.warning("TAVILY_API_KEY not set; web search disabled for all modes")

        retention = int(config["memory"].get("void_retention_days", 30))
        expired = self.memory.purge_expired_tombstones(retention)
        if expired:
            log.info("Purged %d expired tombstones", len(expired))
            self.flog.log_purge_expired(len(expired))

    async def setup_hook(self):
        await self.add_cog(LoreCog(self))

    # ------------------------------------------------------------------
    # Channel + role checks
    # ------------------------------------------------------------------
    def _is_scanned(self, channel) -> bool:
        name = normalize_channel_name(getattr(channel, "name", "") or "")
        if name in self._ignored_set:
            return False
        if self._scan_set:
            return name in self._scan_set
        return True

    def _get_min_length(self, message: discord.Message) -> int:
        has_doc = any(_doc_attachment_type(a) for a in message.attachments)
        if has_doc:
            return 0
        has_image = any(
            a.content_type and a.content_type.startswith("image/")
            for a in message.attachments
        )
        if has_image:
            return self.cfg["bot"].get("min_image_message_length", 10)
        return self.cfg["bot"].get("min_message_length", 0)

    # OOC pre-filter: cheap deterministic check before touching the LLM.
    # Matches common out-of-character notation used in Discord roleplays.
    _OOC_PREFIX_RE = re.compile(
        r"^\s*(\[ooc\]|\(ooc\)|/ooc|\(\(|\[\[|\[jk\]|\[meta\])",
        re.IGNORECASE,
    )
    _OOC_WRAP_RE = re.compile(r"^\s*\(\(.*\)\)\s*$|^\s*\[\[.*\]\]\s*$", re.DOTALL)
    _REACTION_ONLY_RE = re.compile(
        r"^\s*(lol|lmao|lmfao|haha|hehe|xd|x3|brb|afk|omg|gg|rip|f|nice|same|wow|yep|nope|ok|okay|sure|yeah|yea|nah|hmm|hm|ah|oh|oof)\s*[!?.]*\s*$",
        re.IGNORECASE,
    )

    def _is_obviously_ooc(self, text: str) -> bool:
        """Return True if the raw message text is obviously out-of-character."""
        if self._OOC_PREFIX_RE.match(text):
            return True
        if self._OOC_WRAP_RE.match(text):
            return True
        if self._REACTION_ONLY_RE.match(text):
            return True
        return False

    def _is_admin_channel(self, channel) -> bool:
        name = normalize_channel_name(getattr(channel, "name", "") or "")
        return name in self._admin_set

    def _help_text(self, message: discord.Message) -> str:
        prefix = self.prefix
        is_admin = self._is_admin_user(message)
        is_gm = self._is_gm_user(message.author)
        lines = [
            f"**DNC Lore Bot** — `{prefix} <question>` to query the archive",
            "",
            "**Public commands:**",
            f"`{prefix} <question>` — Query the lore archive",
            f"`{prefix} chat <message>` — Chat with the bot",
            f"`{prefix} gm <message-link>` — GM adjudication (GM/admin only)",
            f"`{prefix} war help` — War GM mode (run inside a war's thread)",
            f"`{prefix} year` — Show current in-game year",
            f"`{prefix} optout` / `{prefix} optin` — Manage your opt-out",
            f"`{prefix} whoami` — Show your role/permission status",
            f"`{prefix} map` — Render the player ownership map",
            f"`{prefix} mapfaction` — Render the faction membership map",
            f"`{prefix} mapzoom <TAG>` — Render a focused country map",
            f"`{prefix} maplist [TAG]` — List province ownership",
            f"`{prefix} help` — Show this message",
        ]
        if is_admin or is_gm:
            lines += [
                "",
                "**Admin commands** (admin channel only):",
                f"`{prefix} chatuc <message>` — Unhinged chat mode",
                f"`{prefix} ingest [#channel] <N|date-range>` — Backfill messages",
                f"`{prefix} void <@user|link|ID>` — Void memory(ies)",
                f"`{prefix} unvoid <ID>` — Restore voided memory group",
                f"`{prefix} yearset <year>` — Set in-game year",
                f"`{prefix} yearroll` — Increment year and announce",
                f"`{prefix} purge year <year>` — Void all memories from a year",
                f"`{prefix} export` — Export all memories to files",
                f"`{prefix} stats [reset]` — Token/message statistics",
                f"`{prefix} channels` — Diagnostic channel listing",
                f"`{prefix} reloadprompts` — Hot-reload prompt files",
                f"`{prefix} chatban <@user|ID>` — Block user from chat commands",
                f"`{prefix} chatunban <@user|ID>` — Unblock user from chat commands",
            ]
        return "\n".join(lines)

    def _is_admin_user(self, message: discord.Message) -> bool:
        """True if user has Manage Server perm OR has any admin_roles role."""
        perms = message.channel.permissions_for(message.author)
        if perms.manage_guild:
            return True
        if not isinstance(message.author, discord.Member):
            return False
        if not self._admin_roles:
            return False
        member_roles = {r.name.lower() for r in message.author.roles}
        return bool(member_roles & self._admin_roles)

    def _is_gm_user(self, member: discord.abc.User) -> bool:
        """True if user has any gm_roles role.
        If gm_roles is empty in config, falls back to admin check.
        """
        if not isinstance(member, discord.Member):
            return False
        if not self._gm_roles:
            # No GM roles configured; fall back to admin
            perms = member.guild_permissions
            if perms.manage_guild:
                return True
            if not self._admin_roles:
                return False
            member_role_names = {r.name.lower() for r in member.roles}
            return bool(member_role_names & self._admin_roles)
        member_role_names = {r.name.lower() for r in member.roles}
        return bool(member_role_names & self._gm_roles)

    def _is_war_gm(self, member: discord.abc.User) -> bool:
        """True if user may run GM-only war commands and revise war rulings.

        Uses war.roles if configured; otherwise falls back to GM/admin checks.
        """
        if not isinstance(member, discord.Member):
            return False
        if self._war_roles:
            return bool({r.name.lower() for r in member.roles} & self._war_roles)
        return self._is_gm_user(member)

    def _is_espionage_secret_channel(self, channel) -> bool:
        name = normalize_channel_name(getattr(channel, "name", "") or "")
        return name in self._espionage_secret_set

    def _is_espionage_command_channel(self, channel) -> bool:
        if not self._espionage_command_set:
            return True  # no restriction configured
        name = normalize_channel_name(getattr(channel, "name", "") or "")
        return name in self._espionage_command_set

    def _find_member_by_tag(
        self, guild: discord.Guild, tag: str
    ) -> Optional[discord.Member]:
        """Return the guild member whose display name starts with the given TAG."""
        tag_upper = tag.upper()
        for member in guild.members:
            if parse_tag(member.display_name) == tag_upper:
                return member
        return None

    def _is_chatuc_user(self, message: discord.Message) -> bool:
        """True if user has any chatuc_roles role. Falls back to admin_roles if not configured."""
        if not self._chatuc_roles:
            return self._is_admin_user(message)
        if not isinstance(message.author, discord.Member):
            return False
        member_roles = {r.name.lower() for r in message.author.roles}
        return bool(member_roles & self._chatuc_roles)

    def _is_admin_user_by_id(
        self, user: discord.abc.User, guild: Optional[discord.Guild]
    ) -> bool:
        """Check admin status from a raw user object (e.g. in reaction checks)."""
        if guild is None:
            return False
        member = guild.get_member(user.id)
        if member is None:
            return False
        if member.guild_permissions.manage_guild:
            return True
        if not self._admin_roles:
            return False
        return bool({r.name.lower() for r in member.roles} & self._admin_roles)

    def _mode_settings(self, mode: str) -> dict:
        """Returns chat_messages kwargs for a given mode, read from config.models.modes."""
        m_cfg = self.cfg.get("models", {})
        mode_cfg = m_cfg.get("modes", {}).get(mode, {})
        default_model = m_cfg.get("defaults", {}).get("model")
        return {
            "model":           mode_cfg.get("model") or default_model or None,
            "temperature":     mode_cfg.get("temperature"),
            "top_p":           mode_cfg.get("top_p"),
            "top_k":           mode_cfg.get("top_k"),
            "max_tokens":      mode_cfg.get("max_tokens"),
            "thinking_budget": mode_cfg.get("thinking_budget"),
        }

    def _gm_retrieval_cfg(self) -> dict:
        """Read gm.retrieval config block with safe defaults."""
        r = (self.cfg.get("gm") or {}).get("retrieval") or {}
        return {
            "author_recency_top_n": int(r.get("author_recency_top_n", 30)),
            "author_semantic_top_k": int(r.get("author_semantic_top_k", 15)),
            "wider_top_k": int(r.get("wider_top_k", 15)),
            "include_author_in_wider": bool(r.get("include_author_in_wider", True)),
            "expand_queries": bool(r.get("expand_queries", True)),
            "expansion_model": r.get("expansion_model") or "mistralai/ministral-14b-2512",
            "expansion_max_queries": int(r.get("expansion_max_queries", 4)),
            "action_type_boost": bool(r.get("action_type_boost", True)),
        }

    async def _gm_expand_queries(
        self, action_text: str, max_queries: int = 4, model: str | None = None,
    ) -> tuple[list[str], Optional[str]]:
        """Generate search queries and infer an action_type for GM retrieval.

        Returns (queries, action_type_or_None). On failure returns ([], None).
        """
        if not action_text.strip():
            return [], None
        sys_prompt = (
            "You are a search-query generator for a roleplay archive. "
            "Given an action being adjudicated, output:\n"
            "Line 1: ACTION_TYPE: <internal|military|economic|diplomatic|cultural|other>\n"
            f"Lines 2-{max_queries + 1}: short search queries (one per line, no numbering, "
            "no questions) that would surface relevant prior context — named entities, "
            "prerequisites the action assumes, claims being built upon.\n"
            "Each query is a short noun phrase or assertion. Output only those lines."
        )
        try:
            kwargs = {"temperature": 0.2}
            if model:
                kwargs["model"] = model
            raw, usage = await self.llm.chat(
                sys_prompt,
                _trim_context_text(action_text, _EMBED_INPUT_CHARS),
                **kwargs,
            )
            self.state.add_usage("archivist", usage)
        except Exception:
            log.exception("GM query expansion failed; continuing with original action text only")
            return [], None

        action_type: Optional[str] = None
        queries: list[str] = []
        for line in (raw or "").splitlines():
            line = line.strip()
            if not line:
                continue
            if line.upper().startswith("ACTION_TYPE:"):
                v = line.split(":", 1)[1].strip().lower().split()[:1]
                if v and v[0] in _ARCHIVIST_VALID_ACTION_TYPES:
                    action_type = v[0]
                continue
            # Strip common list prefixes the model might add
            cleaned = re.sub(r"^[\-\*\d\.\)\s]+", "", line).strip()
            if cleaned and len(cleaned) <= 200:
                queries.append(cleaned)
            if len(queries) >= max_queries:
                break
        return queries, action_type

    def _parse_date(self, date_str: str) -> Optional[datetime]:
        """Attempt to parse common date formats into UTC datetimes."""
        formats = ["%m-%d-%Y", "%Y-%m-%d", "%m/%d/%Y", "%Y/%m/%d"]
        for fmt in formats:
            try:
                dt = datetime.strptime(date_str, fmt)
                return dt.replace(tzinfo=timezone.utc)
            except ValueError:
                pass
        return None

    # ------------------------------------------------------------------
    async def on_ready(self):
        self.state.mark_started()
        log.info("Logged in as %s — %d memories, year=%d",
                 self.user, self.memory.count(), self.state.current_year)
        if self.cfg["bot"].get("year_rollover", {}).get("enabled"):
            if not self._rollover_task.is_running():
                self._rollover_task.start()
        if self.cfg.get("spatial_mapping", {}).get("enabled"):
            if not self._map_render_task.is_running():
                self._map_render_task.start()
        if self.cfg.get("spatial_mapping", {}).get("prewarm_on_startup", True):
            asyncio.create_task(self._prewarm_map_cache())

    async def on_member_remove(self, member: discord.Member):
        if self.optouts.remove(str(member.id)):
            log.info("Auto-pruned opt-out for %s (%s)", member.display_name, member.id)

    async def on_member_update(self, before: discord.Member, after: discord.Member):
        if before.nick == after.nick:
            return
        if not self.cfg.get("spatial_mapping", {}).get("enabled"):
            return
        old_tag = parse_tag(before.nick or before.name)
        new_tag = parse_tag(after.nick or after.name)
        if old_tag == new_tag:
            return
        if old_tag:
            self.map_store.clear_occupation(old_tag)
            log.info("Map: vacated %s (nick change by %s)", old_tag, after.display_name)
        if new_tag:
            self.map_store.set_occupation(new_tag, str(after.id), after.display_name)
            log.info("Map: assigned %s → %s", new_tag, after.display_name)
        await self._render_and_post_map()

    async def on_raw_reaction_add(self, payload: discord.RawReactionActionEvent):
        if str(payload.emoji) != "🎲":
            return
        if self.user and payload.user_id == self.user.id:
            return

        msg_id = str(payload.message_id)

        # War ruling reroll: original move submitter or a war GM may reroll.
        war_cache = self._war_ruling_cache.get(msg_id)
        if war_cache is not None:
            war = self.war_store.get_active_war(war_cache["thread_id"])
            if not war or war.get("status") != "active":
                return
            guild = self.get_guild(payload.guild_id)
            if guild is None:
                return
            member = guild.get_member(payload.user_id)
            allowed = (
                (war_cache.get("author_id")
                 and str(payload.user_id) == war_cache["author_id"])
                or (member is not None and self._is_war_gm(member))
            )
            if not allowed:
                return
            thread = guild.get_thread(int(war_cache["thread_id"]))
            if thread is None:
                try:
                    thread = await self.fetch_channel(int(war_cache["thread_id"]))
                except (discord.NotFound, discord.Forbidden, discord.HTTPException):
                    return
            await self._produce_war_ruling(
                thread=thread, tid=war_cache["thread_id"], invocation=None,
                actor_tag=war_cache["actor_tag"], actor_side=war_cache["actor_side"],
                is_npc=war_cache["is_npc"], move_text=war_cache["move_text"],
                source_message_id=war_cache["source_message_id"],
                author_id=war_cache.get("author_id", ""),
                author_name=war_cache.get("author_name", ""),
                allow_npc_response=False, guidance=war_cache.get("guidance"),
                existing_move_seq=war_cache["move_seq"],
                old_ruling_seq=war_cache["ruling_seq"],
                old_ruling_msg_ids=war_cache.get("ruling_msg_ids"),
                is_reroll=True,
            )
            return

        cache = self._ruling_msg_cache.get(msg_id)
        if cache is None:
            return

        # Only the original action author may reroll
        if payload.user_id != cache["action_author_id"]:
            return

        guild = self.get_guild(payload.guild_id)
        if guild is None:
            return

        target_channel = guild.get_channel(cache["target_channel_id"])
        if target_channel is None:
            return
        try:
            target_msg = await target_channel.fetch_message(cache["target_msg_id"])
        except (discord.NotFound, discord.Forbidden):
            return

        # Try to fetch original invocation message for channel/author context
        invocation_msg: Optional[discord.Message] = None
        inv_ch = guild.get_channel(cache["invocation_channel_id"])
        if inv_ch:
            try:
                invocation_msg = await inv_ch.fetch_message(cache["invocation_msg_id"])
            except (discord.NotFound, discord.Forbidden):
                pass
        if invocation_msg is None:
            reaction_ch = guild.get_channel(payload.channel_id)
            if reaction_ch is None:
                return
            try:
                invocation_msg = await reaction_ch.fetch_message(payload.message_id)
            except (discord.NotFound, discord.Forbidden):
                return

        output_channel = None
        if self._gm_output_channel_name:
            output_channel = discord.utils.find(
                lambda c: normalize_channel_name(c.name) == self._gm_output_channel_name,
                guild.text_channels,
            )

        await self._produce_gm_ruling(
            invocation=invocation_msg,
            target_msg=target_msg,
            chain_back=[],
            revision_request=None,
            output_channel=output_channel,
            focus_question=cache.get("focus_question"),
            is_reroll=True,
        )

    # ------------------------------------------------------------------
    @tasks.loop(hours=1)
    async def _rollover_task(self):
        cfg = self.cfg["bot"].get("year_rollover", {})
        if not cfg.get("enabled"):
            return
        interval = int(cfg.get("interval_days", 7))
        if not year_scheduler.is_overdue(self.state, interval):
            return
        await year_scheduler.prompt_for_overdue_rollover(
            state=self.state,
            client=self,
            year_check_channel=cfg.get("year_check_channel", ""),
            on_announce=self._announce_year_roll,
        )

    @_rollover_task.before_loop
    async def _before_rollover(self):
        await self.wait_until_ready()

    @tasks.loop(hours=24)
    async def _map_render_task(self):
        if not self.cfg.get("spatial_mapping", {}).get("enabled"):
            return
        for guild in self.guilds:
            self.map_store.rebuild(map_sched.scan_members(guild.members))
        await self._render_and_post_map()

    @_map_render_task.before_loop
    async def _before_map_render(self):
        await self.wait_until_ready()
        for guild in self.guilds:
            self.map_store.rebuild(map_sched.scan_members(guild.members))
        log.info("Map: initial scan complete — %d nations occupied", len(self.map_store.get_occupation()))

    async def _announce_year_roll(self, year: int):
        cfg = self.cfg["bot"].get("year_rollover", {})
        template = cfg.get("announcement_template", "📜 The year turns. It is now {year}.")
        body = template.format(year=year)
        targets = {normalize_channel_name(c)
                   for c in cfg.get("announcement_channels", []) or []}
        if not targets:
            return
        for guild in self.guilds:
            for ch in guild.text_channels:
                if normalize_channel_name(ch.name) in targets:
                    try:
                        await ch.send(body)
                    except discord.HTTPException:
                        log.exception("Failed announcement in #%s", ch.name)

    # ------------------------------------------------------------------
    # Message routing
    # ------------------------------------------------------------------
    async def on_message(self, message: discord.Message):
        if message.author.bot or not message.guild:
            return

        self.state.bump("messages_seen")
        content = message.content.strip()
        is_opted_out = self.optouts.is_opted_out(str(message.author.id))

        # Espionage interception — fires for any post in a secret channel,
        # independently of lore archiving and opt-out status.
        if (self._espionage_cfg.get("enabled")
                and self._is_espionage_secret_channel(message.channel)):
            asyncio.ensure_future(self._check_espionage(message))

        # 1. Reply-to-bot path: continuation turn
        if message.reference and message.reference.message_id:
            try:
                parent = await message.channel.fetch_message(message.reference.message_id)
                if parent.author.id == self.user.id:
                    # If it's a reply to us, and NOT starting with the prefix, handle as continuation
                    if not message.content.startswith(self.prefix):
                        if is_opted_out:
                            await message.reply(
                                f"You're currently opted out of the lore archive. "
                                f"Use `{self.prefix} optin` to re-enable access."
                            )
                            return
                        await self._handle_reply_to_bot(message, parent)
                        return
            except (discord.NotFound, discord.Forbidden, discord.HTTPException):
                pass  # fall through

        # 2. Process commands (this handles the !DNC prefix logic via commands.Bot)
        # cog_check in LoreCog enforces opt-out for all commands except optout/optin
        ctx = await self.get_context(message)
        if ctx.valid:
            await self.invoke(ctx)
            return

        # 3. Fallback: If it starts with prefix but isn't a known command, treat as Lore Query
        if content.startswith(self.prefix):
            query_text = content[len(self.prefix):].strip()
            if query_text:
                if is_opted_out:
                    await message.reply(
                        f"You're currently opted out of the lore archive. "
                        f"Use `{self.prefix} optin` to re-enable access."
                    )
                    return
                await self._cmd_query(message, query_text)
                return
            else:
                # Just the prefix: show help (always allowed so opted-out users can find optin)
                await message.reply(self._help_text(message))
                return

        # 4. Live ingestion path (non-command messages)
        # War threads are isolated from the main archive while the war is live;
        # only the final chronicle (on `war commit`) is written to lore.
        if (isinstance(message.channel, discord.Thread)
                and self.war_store.get_active_war(str(message.channel.id))):
            self.state.bump("messages_filtered_local"); return
        if not self._is_scanned(message.channel):
            self.state.bump("messages_filtered_local"); return
        if len(content) < self._get_min_length(message):
            self.state.bump("messages_filtered_local"); return
        if self.optouts.is_opted_out(str(message.author.id)):
            self.state.bump("messages_filtered_local"); return
        if self._is_obviously_ooc(content):
            self.state.bump("messages_filtered_local"); return

        try:
            await self._queue_for_ingestion(message)
        except Exception:
            log.exception("Memory ingestion failed for %s", message.id)

    # ------------------------------------------------------------------
    # Reply handling — mode detection + routing
    # ------------------------------------------------------------------
    async def _handle_reply_to_bot(self, message: discord.Message,
                                    parent: discord.Message):
        # War ruling revision: a GM replies to a war ruling asking for changes.
        war_cache = self._war_ruling_cache.get(str(parent.id))
        if war_cache is not None:
            if not self._is_war_gm(message.author):
                return  # silent — only war GMs revise war rulings
            if not isinstance(message.channel, discord.Thread):
                return
            await self._produce_war_ruling(
                thread=message.channel, tid=war_cache["thread_id"],
                invocation=message,
                actor_tag=war_cache["actor_tag"], actor_side=war_cache["actor_side"],
                is_npc=war_cache["is_npc"], move_text=war_cache["move_text"],
                source_message_id=war_cache["source_message_id"],
                author_id=war_cache.get("author_id", ""),
                author_name=war_cache.get("author_name", ""),
                allow_npc_response=False, guidance=war_cache.get("guidance"),
                revision_request=message.content,
                existing_move_seq=war_cache["move_seq"],
                old_ruling_seq=war_cache["ruling_seq"],
                old_ruling_msg_ids=war_cache.get("ruling_msg_ids"),
            )
            return

        # Walk the reply chain to find the conversation root and the mode.
        chain_back, root = await reply_chain.walk_reply_chain(
            message=message,
            bot_user_id=self.user.id,
            max_depth=self._max_chain_depth,
        )

        # Determine mode from the root invocation
        mode = "query"
        if root is not None:
            mode = reply_chain.detect_mode_from_root(root.content or "", self.prefix)

        # Mode-specific access control
        if mode == "chatuc":
            # Chatuc threads are admin-only end-to-end
            if not self._is_admin_user(message):
                return  # silent
        elif mode == "gm":
            # GM ruling replies are gm_roles-only
            if not self._is_gm_user(message.author):
                return  # silent
        # chat and query replies: anyone can continue

        # Route to the appropriate handler
        if mode == "chat":
            await self._reply_chat(message, chain_back, mode="chat")
        elif mode == "chatuc":
            await self._reply_chat(message, chain_back, mode="chatuc")
        elif mode == "gm":
            await self._reply_gm_revision(message, chain_back, root)
        else:
            await self._reply_query(message, chain_back, root)

    # ------------------------------------------------------------------
    # CHAT and CHATUC (initial invocation)

    # ------------------------------------------------------------------
    async def _cmd_chat(self, message: discord.Message, body: str):
        if not body:
            await message.reply(f"Usage: `{self.prefix} chat <something to say>`")
            return
        await self._do_chat(message, body, prompt_name="chat", mode_key="chat",
                            stat_count="chats_answered", stat_usage="chat",
                            chain_back=[])

    async def _cmd_chatuc(self, message: discord.Message, body: str):
        if not body:
            await message.reply(f"Usage: `{self.prefix} chatuc <message>`")
            return
        await self._do_chat(message, body, prompt_name="chat_unhinged", mode_key="chatuc",
                            stat_count="chatuc_answered", stat_usage="chatuc",
                            chain_back=[])

    async def _do_chat(self, message: discord.Message, body: str,
                       prompt_name: str, mode_key: str, stat_count: str, stat_usage: str,
                       chain_back: List[discord.Message]):
        if self.chat_blacklist.is_blacklisted(str(message.author.id)):
            await message.reply("You are not permitted to use chat commands.")
            return

        try:
            system_prompt = self.prompts.get(prompt_name)
        except KeyError:
            await message.reply(f"⚠️ Prompt `{prompt_name}` not configured.")
            return

        async with message.channel.typing():
            try:
                msgs = reply_chain.build_messages_from_chain(
                    chain_back=chain_back,
                    current_message=message,
                    bot_user_id=self.user.id,
                    system_prompt=system_prompt,
                    max_message_chars=self._max_history_chars,
                )
                _cs = self._mode_settings(mode_key)
                tavily_enabled = (
                    self.tavily is not None
                    and self.cfg.get("tavily", {})
                        .get("enabled_modes", {})
                        .get(mode_key, False)
                )
                print(f"{_C_SEND}[{mode_key.upper():<8}→] model={_cs['model'] or 'default'}"
                      f"  temp={_cs['temperature']}  thinking={_fmt_think(_cs['thinking_budget'])}"
                      f"  search={'on' if tavily_enabled else 'off'}{_C_RST}",
                      flush=True)
                answer, usage = await self.llm.chat_messages(
                    msgs, **_cs,
                    tools=[self.tavily.tool_definition()] if tavily_enabled else None,
                    tool_executor=self.tavily.execute if tavily_enabled else None,
                )
                self.state.add_usage(stat_usage, usage)
                self.state.bump(stat_count)
                await self._send_chunked_reply(message, answer)
            except Exception as e:
                log.exception("%s failed", prompt_name)
                await message.reply(f"⚠️ {prompt_name} failed: `{e}`")

    # Reply continuation for chat / chatuc
    async def _reply_chat(self, message: discord.Message,
                          chain_back: List[discord.Message],
                          mode: str):
        if mode == "chatuc":
            prompt_name, stat_count, stat_usage, mode_key = (
                "chat_unhinged", "chatuc_answered", "chatuc", "chatuc")
        else:
            prompt_name, stat_count, stat_usage, mode_key = (
                "chat", "chats_answered", "chat", "chat")
        await self._do_chat(message, message.content, prompt_name=prompt_name,
                            mode_key=mode_key, stat_count=stat_count,
                            stat_usage=stat_usage, chain_back=chain_back)

    # ------------------------------------------------------------------
    # LORE QUERIES
    # ------------------------------------------------------------------
    async def _cmd_query(self, message: discord.Message, question: str):
        await self._do_query(message, question, chain_back=[])

    async def _reply_query(self, message: discord.Message,
                           chain_back: List[discord.Message],
                           root: Optional[discord.Message]):
        # Use the current message as the new question; fresh search per turn.
        await self._do_query(message, message.content, chain_back=chain_back)

    async def _do_query(self, message: discord.Message, question: str,
                        chain_back: List[discord.Message]):
        async with message.channel.typing():
            try:
                print(f"{_C_SEND}[QUERY EMB→] {question[:120]!r}{_C_RST}", flush=True)
                q_embed, embed_usage = await self.llm.embed(question)
                self.state.add_usage("embedding", embed_usage)
                hits = self.memory.search(q_embed, top_k=self.cfg["memory"]["top_k"])
                print(f"{_C_RECV}[QUERY EMB←] dims={len(q_embed)}  hits={len(hits)}{_C_RST}", flush=True)
                if not hits:
                    await message.reply("*The archives are empty.*")
                    return

                archive_lines = []
                for i, h in enumerate(hits):
                    meta = h["metadata"]
                    full_text = meta.get("full_message_text", "") or meta.get(
                        "original_excerpt", "")
                    a_char = meta.get("author_name", "?")
                    a_acct = meta.get("author_discord_name", "")
                    a_label = (f"{a_char} (@{a_acct})"
                               if a_acct and a_acct != a_char else a_char)
                    archive_lines.append(
                        f"[Excerpt {i+1}] (#{meta.get('channel', '?')}, "
                        f"year {meta.get('year', '?')}, "
                        f"{str(meta.get('timestamp', ''))[:10]}, "
                        f"by {a_label})\n"
                        f"Summary: {h['text']}\n"
                        f"Full post: {full_text}"
                    )
                archive = "\n\n".join(archive_lines)

                system_prompt = self.prompts.get("lore_query")
                # Build messages: system + chain history + final user query
                msgs = [{"role": "system", "content": system_prompt}]
                # For chain history, just include it as conversational context
                for cm in chain_back:
                    if cm.author.id == self.user.id:
                        c = (cm.content or "").strip()
                        if len(c) > self._max_history_chars:
                            c = c[:self._max_history_chars] + "…"
                        msgs.append({"role": "assistant", "content": c})
                    else:
                        c = (cm.content or "").strip()
                        if c.lower().startswith("!dnc"):
                            parts = c.split(maxsplit=1)
                            c = parts[1] if len(parts) > 1 else ""
                        if len(c) > self._max_history_chars:
                            c = c[:self._max_history_chars] + "…"
                        cm_char = cm.author.display_name
                        cm_acct = cm.author.name
                        cm_label = (f"{cm_char} (@{cm_acct})"
                                    if cm_char != cm_acct else cm_char)
                        msgs.append({"role": "user",
                                     "content": f"[{cm_label}]: {c}"})
                msgs.append({"role": "user",
                             "content": f"Archive excerpts:\n{archive}\n\nQuestion: {question}"})

                _qs = self._mode_settings("query")
                print(f"{_C_SEND}[QUERY    →] model={_qs['model'] or 'default'}"
                      f"  temp={_qs['temperature']}  thinking={_fmt_think(_qs['thinking_budget'])}{_C_RST}",
                      flush=True)
                answer, usage = await self.llm.chat_messages(
                    msgs, **_qs
                )
                self.state.add_usage("query", usage)
                self.state.bump("queries_answered")
                await self._send_chunked_reply(message, answer)
            except Exception as e:
                log.exception("Query failed")
                await message.reply(f"⚠️ Query failed: `{e}`")

    # ------------------------------------------------------------------
    # GM mode
    # ------------------------------------------------------------------
    async def _cmd_gm(self, message: discord.Message, arg: str):
        if not self._is_gm_user(message.author):
            await message.reply("⚠️ You are not permitted to request GM rulings.")
            return

        arg = arg.strip()
        if not arg:
            await message.reply(
                f"Usage: `{self.prefix} gm <message-link> [question]` — paste a Discord "
                f"message link to the action you want adjudicated. Optionally add a "
                f"specific question, e.g. `how does this affect GDP?`"
            )
            return

        link = MSG_LINK_RE.match(arg)
        if not link:
            await message.reply(
                "⚠️ Please provide a Discord message link "
                "(right-click on the message → Copy Message Link)."
            )
            return

        focus_question = arg[link.end():].strip() or None
        guild_id, channel_id, msg_id = link.group(1), link.group(2), link.group(3)
        if str(message.guild.id) != guild_id:
            await message.reply("⚠️ That message is in a different server.")
            return

        target_channel = message.guild.get_channel(int(channel_id))
        if target_channel is None:
            await message.reply("⚠️ I can't see that channel.")
            return

        if self._gm_channel_set:
            chan_norm = normalize_channel_name(getattr(message.channel, "name", "") or "")
            if chan_norm not in self._gm_channel_set:
                await message.reply(
                    f"⚠️ `{self.prefix} gm` must be used in a designated GM channel."
                )
                return

        try:
            target_msg = await target_channel.fetch_message(int(msg_id))
        except discord.NotFound:
            await message.reply("⚠️ Couldn't find that message.")
            return
        except discord.Forbidden:
            await message.reply("⚠️ I don't have permission to read that channel.")
            return

        output_channel = None
        if self._gm_output_channel_name:
            output_channel = discord.utils.find(
                lambda c: normalize_channel_name(c.name) == self._gm_output_channel_name,
                message.guild.text_channels,
            )
            if output_channel is None:
                await message.reply(
                    f"⚠️ GM output channel `{self._gm_output_channel_name}` not found — "
                    f"check `gm_output_channel` in config and bot channel permissions."
                )
                return

        await self._produce_gm_ruling(invocation=message, target_msg=target_msg,
                                       chain_back=[], revision_request=None,
                                       output_channel=output_channel,
                                       focus_question=focus_question)

    async def _reply_gm_revision(self, message: discord.Message,
                                  chain_back: List[discord.Message],
                                  root: Optional[discord.Message]):
        """A GM has replied to a ruling asking for revision."""
        if root is None:
            await message.reply("⚠️ Couldn't find the original ruling context.")
            return
        link_match = MSG_LINK_RE.search(root.content or "")
        if not link_match:
            await message.reply("⚠️ Couldn't find the ruled-on message link in the root.")
            return
        try:
            target_channel = message.guild.get_channel(int(link_match.group(2)))
            target_msg = await target_channel.fetch_message(int(link_match.group(3)))
        except (discord.NotFound, discord.Forbidden, AttributeError):
            await message.reply("⚠️ Couldn't re-fetch the original action message.")
            return

        output_channel = None
        if self._gm_output_channel_name:
            output_channel = discord.utils.find(
                lambda c: normalize_channel_name(c.name) == self._gm_output_channel_name,
                message.guild.text_channels,
            )

        await self._produce_gm_ruling(
            invocation=message, target_msg=target_msg,
            chain_back=chain_back,
            revision_request=message.content,
            output_channel=output_channel,
        )
        self.state.bump("gm_rulings_revised")

    async def _produce_gm_ruling(self, invocation: discord.Message,
                                  target_msg: discord.Message,
                                  chain_back: List[discord.Message],
                                  revision_request: Optional[str],
                                  output_channel: Optional[discord.TextChannel] = None,
                                  focus_question: Optional[str] = None,
                                  is_reroll: bool = False):
        gm_key = (
            f"reroll:{invocation.id}:{target_msg.id}" if is_reroll else
            f"revision:{invocation.id}:{target_msg.id}" if revision_request else
            f"gm:{invocation.id}:{target_msg.id}"
        )
        if gm_key in self._gm_inflight:
            log.info("Skipping duplicate GM request %s", gm_key)
            return
        self._gm_inflight.add(gm_key)
        dest = output_channel or invocation.channel
        try:
            async with dest.typing():
                gm_rcfg = self._gm_retrieval_cfg()
                target_author_id = str(target_msg.author.id)
                action_iso = target_msg.created_at.isoformat()
                is_replacement = bool(revision_request) or is_reroll

                # --- Determine ruling ID up front so it can be injected into the prompt ---
                ruling_id: Optional[str] = None
                if is_replacement:
                    # Revisions/rerolls keep the original ID
                    existing = self.memory.get_rulings_by_target_message(str(target_msg.id))
                    for old in existing:
                        rid = (old.get("metadata") or {}).get("ruling_id") or ""
                        if rid:
                            ruling_id = rid
                            break
                if not ruling_id:
                    seq = self.state.next_ruling_id(self.state.current_year)
                    ruling_id = f"GM-{self.state.current_year}-{seq:04d}"

                # --- Collect the full action: may span multiple consecutive posts ---
                chain_msgs = await self._gather_action_chain(target_msg)
                if len(chain_msgs) > 1:
                    action_text = "\n\n".join(
                        await self._message_text_with_documents(m)
                        for m in chain_msgs
                        if (m.content or "").strip() or any(_doc_attachment_type(a) for a in m.attachments)
                    )
                else:
                    action_text = await self._message_text_with_documents(target_msg)

                # --- Optional LLM query expansion (entities, prerequisites, action_type) ---
                expansion_queries: list[str] = []
                inferred_action_type: Optional[str] = None
                if action_text and gm_rcfg["expand_queries"]:
                    expansion_queries, inferred_action_type = await self._gm_expand_queries(
                        action_text,
                        max_queries=gm_rcfg["expansion_max_queries"],
                        model=gm_rcfg["expansion_model"],
                    )
                    if expansion_queries or inferred_action_type:
                        print(f"{_C_DIM}[GM EXP    ] action_type={inferred_action_type or '-'} "
                              f"queries={expansion_queries}{_C_RST}", flush=True)

                # --- Embed the action text (and any expansion queries) ---
                action_embed: Optional[list] = None
                expansion_embeds: list[list] = []
                if action_text:
                    try:
                        search_text = _trim_context_text(action_text, _EMBED_INPUT_CHARS)
                        print(f"{_C_SEND}[GM EMB   →] action text {len(search_text)} chars{_C_RST}", flush=True)
                        action_embed, embed_usage = await self.llm.embed(search_text)
                        self.state.add_usage("embedding", embed_usage)
                        print(f"{_C_RECV}[GM EMB   ←] dims={len(action_embed)}{_C_RST}", flush=True)
                    except Exception:
                        log.exception("Action embedding failed; wider/semantic-prior retrieval disabled")
                for q in expansion_queries:
                    try:
                        e, u = await self.llm.embed(q)
                        self.state.add_usage("embedding", u)
                        expansion_embeds.append(e)
                    except Exception:
                        log.exception("Expansion query embedding failed for %r; skipping", q[:80])

                all_query_embeds: list[list] = []
                if action_embed:
                    all_query_embeds.append(action_embed)
                all_query_embeds.extend(expansion_embeds)

                # --- Retrieve author priors: recency pass + semantic pass, deduped ---
                recency_priors = self.memory.get_by_author_before(target_author_id, action_iso)
                recency_priors.sort(key=lambda m: m["metadata"].get("timestamp", ""), reverse=True)
                recency_priors = recency_priors[:gm_rcfg["author_recency_top_n"]]
                recency_keys = {
                    p["metadata"].get("source_message_id") or p["id"]
                    for p in recency_priors
                }
                for p in recency_priors:
                    p["_source"] = "recency"

                semantic_prior_hits: list[dict] = []
                if all_query_embeds:
                    seen_ids: set[str] = set()
                    best_by_id: dict[str, dict] = {}
                    for emb in all_query_embeds:
                        try:
                            hits = self.memory.search(
                                emb,
                                top_k=gm_rcfg["author_semantic_top_k"],
                                where={"author_id": target_author_id},
                            )
                        except Exception:
                            log.exception("Author-semantic search failed for one query; skipping")
                            continue
                        for h in hits:
                            ts = (h["metadata"].get("timestamp") or "")
                            if ts >= action_iso:
                                continue
                            key = h["metadata"].get("source_message_id") or h["id"]
                            if key in recency_keys:
                                continue
                            prev = best_by_id.get(key)
                            if (prev is None) or (h.get("distance", 1.0) < prev.get("distance", 1.0)):
                                best_by_id[key] = h
                            seen_ids.add(key)
                    for h in best_by_id.values():
                        h["_source"] = "semantic"
                    semantic_prior_hits = list(best_by_id.values())

                merged_priors = recency_priors + semantic_prior_hits
                merged_priors.sort(key=lambda m: m["metadata"].get("timestamp", ""), reverse=True)

                if merged_priors:
                    prior_lines = []
                    remaining_prior_chars = _GM_PRIOR_CONTEXT_CHARS
                    for p in merged_priors:
                        meta = p["metadata"]
                        full = meta.get("full_message_text", "") or meta.get("original_excerpt", "")
                        full = _trim_context_text(
                            full, min(_GM_CONTEXT_ITEM_CHARS, remaining_prior_chars),
                        )
                        remaining_prior_chars = max(0, remaining_prior_chars - len(full))
                        tag = "[semantic match]" if p.get("_source") == "semantic" else "[recency]"
                        prior_lines.append(
                            f"{tag} [{str(meta.get('timestamp', ''))[:10]}, "
                            f"#{meta.get('channel', '?')}, year {meta.get('year', '?')}]\n"
                            f"Summary: {p['text']}\n"
                            f"Full post: {full}"
                        )
                        if remaining_prior_chars <= 0:
                            prior_lines.append("[Additional prior posts omitted to keep GM context within budget.]")
                            break
                    priors_block = "\n\n".join(prior_lines)
                else:
                    priors_block = ("(No prior posts archived for this author. "
                                    "The author has not established context.)")

                # --- Wider semantic context: no longer excludes author; deduped vs priors ---
                wider_block = ""
                priors_keys = {
                    p["metadata"].get("source_message_id") or p["id"]
                    for p in merged_priors
                }
                if all_query_embeds:
                    best_by_id: dict[str, dict] = {}
                    for emb in all_query_embeds:
                        try:
                            hits = self.memory.search(emb, top_k=gm_rcfg["wider_top_k"])
                        except Exception:
                            log.exception("Wider search failed for one query; skipping")
                            continue
                        for h in hits:
                            if (not gm_rcfg["include_author_in_wider"]
                                    and h["metadata"].get("author_id") == target_author_id):
                                continue
                            key = h["metadata"].get("source_message_id") or h["id"]
                            if key in priors_keys:
                                continue
                            prev = best_by_id.get(key)
                            if (prev is None) or (h.get("distance", 1.0) < prev.get("distance", 1.0)):
                                best_by_id[key] = h

                    # Optional action_type boost: extra pass filtered to inferred type
                    if (gm_rcfg["action_type_boost"] and inferred_action_type
                            and action_embed):
                        try:
                            type_hits = self.memory.search(
                                action_embed,
                                top_k=gm_rcfg["wider_top_k"],
                                where={"action_type": inferred_action_type},
                            )
                            for h in type_hits:
                                if (not gm_rcfg["include_author_in_wider"]
                                        and h["metadata"].get("author_id") == target_author_id):
                                    continue
                                key = h["metadata"].get("source_message_id") or h["id"]
                                if key in priors_keys:
                                    continue
                                prev = best_by_id.get(key)
                                if (prev is None) or (h.get("distance", 1.0) < prev.get("distance", 1.0)):
                                    best_by_id[key] = h
                        except Exception:
                            log.exception("Action-type-boost search failed; continuing")

                    wider_hits = sorted(
                        best_by_id.values(), key=lambda h: h.get("distance", 1.0),
                    )
                    print(f"{_C_RECV}[GM EMB   ←] wider hits={len(wider_hits)} "
                          f"(top_k={gm_rcfg['wider_top_k']}, queries={len(all_query_embeds)}, "
                          f"type_boost={inferred_action_type or '-'}){_C_RST}", flush=True)

                    if wider_hits:
                        wider_lines = []
                        remaining_wider_chars = _GM_WIDER_CONTEXT_CHARS
                        for h in wider_hits:
                            meta = h["metadata"]
                            full = meta.get("full_message_text", "") or meta.get("original_excerpt", "")
                            full = _trim_context_text(
                                full, min(_GM_CONTEXT_ITEM_CHARS, remaining_wider_chars),
                            )
                            remaining_wider_chars = max(0, remaining_wider_chars - len(full))
                            wider_lines.append(
                                f"[{str(meta.get('timestamp', ''))[:10]}, "
                                f"by {meta.get('author_name', '?')}, "
                                f"#{meta.get('channel', '?')}]\n"
                                f"Summary: {h['text']}\n"
                                f"Full post/document excerpt: {full}"
                            )
                            if remaining_wider_chars <= 0:
                                wider_lines.append("[Additional wider-context hits omitted to keep GM context within budget.]")
                                break
                        wider_block = "\n\n".join(wider_lines)

                # --- Build the user-content block ---
                gm_char = target_msg.author.display_name
                gm_acct = target_msg.author.name
                gm_author_line = (
                    f"Author: {gm_char} (Discord: @{gm_acct})"
                    if gm_char != gm_acct
                    else f"Author: {gm_char}"
                )
                chain_note = (
                    f" ({len(chain_msgs)} messages combined)"
                    if len(chain_msgs) > 1 else ""
                )
                user_content_parts = [
                    f"=== RULING ID ===",
                    f"Use this EXACT id in your ruling header: {ruling_id}",
                    f"In-game year: {self.state.current_year}",
                    f"",
                    f"=== ACTION TO ADJUDICATE ===",
                    gm_author_line,
                    f"Posted: {action_iso}",
                    f"Channel: #{target_msg.channel.name}",
                    f"Message link: {target_msg.jump_url}",
                    f"",
                    f"--- Action text{chain_note} ---",
                    action_text,
                    f"",
                    f"=== AUTHOR'S PRIOR POSTS (recency + semantic matches, newest first) ===",
                    priors_block,
                ]
                if wider_block:
                    user_content_parts += [
                        f"",
                        f"=== WIDER ARCHIVE CONTEXT ===",
                        wider_block,
                    ]
                if focus_question:
                    user_content_parts += [
                        f"",
                        f"=== SPECIFIC QUESTION FROM GM ===",
                        focus_question,
                        f"",
                        f"Address this question directly in your ruling.",
                    ]
                if revision_request:
                    user_content_parts += [
                        f"",
                        f"=== GM REVISION REQUEST ===",
                        f"A GM has asked you to revise this ruling. Their request:",
                        revision_request,
                        f"",
                        f"Adjust the ruling per their direction. Re-issue in full.",
                    ]
                user_content = "\n".join(user_content_parts)

                # Build messages: system + history + user content
                system_prompt = self.prompts.get("gm_ruling")
                msgs: List[dict] = [{"role": "system", "content": system_prompt}]
                for cm in chain_back:
                    if cm.author.id == self.user.id:
                        c = (cm.content or "").strip()
                        if len(c) > self._max_history_chars:
                            c = c[:self._max_history_chars] + "…"
                        msgs.append({"role": "assistant", "content": c})
                    else:
                        c = (cm.content or "").strip()
                        if c.lower().startswith("!dnc"):
                            parts = c.split(maxsplit=1)
                            c = parts[1] if len(parts) > 1 else ""
                        if len(c) > self._max_history_chars:
                            c = c[:self._max_history_chars] + "…"
                        cm_char = cm.author.display_name
                        cm_acct = cm.author.name
                        cm_label = (f"{cm_char} (@{cm_acct})"
                                    if cm_char != cm_acct else cm_char)
                        msgs.append({"role": "user",
                                     "content": f"[{cm_label}]: {c}"})
                msgs.append({"role": "user", "content": user_content})

                _gms = self._mode_settings("gm")
                print(f"{_C_SEND}[GM       →] model={_gms['model'] or 'default'}"
                      f"  temp={_gms['temperature']}  thinking={_fmt_think(_gms['thinking_budget'])}{_C_RST}",
                      flush=True)
                ruling, usage = await self.llm.chat_messages(
                    msgs, **_gms
                )
                self.state.add_usage("gm", usage)
                if not is_replacement:
                    self.state.bump("gm_rulings_made")

                # --- Revision/reroll: void old ruling memories and locate old Discord messages ---
                old_ruling_msg_ids: List[str] = []
                old_ruling_channel: Optional[discord.TextChannel] = None
                if is_replacement:
                    old_rulings = self.memory.get_rulings_by_target_message(
                        str(target_msg.id)
                    )
                    for old in old_rulings:
                        meta = old["metadata"]
                        if (not old_ruling_channel
                                and meta.get("ruling_discord_channel_id")):
                            try:
                                ch = invocation.guild.get_channel(
                                    int(meta["ruling_discord_channel_id"])
                                )
                                if ch:
                                    old_ruling_channel = ch
                                    old_ruling_msg_ids = [
                                        mid.strip()
                                        for mid in meta.get(
                                            "ruling_discord_message_ids", ""
                                        ).split(",")
                                        if mid.strip()
                                    ]
                            except (ValueError, AttributeError):
                                pass
                        self.memory.void(
                            memory_ids=[old["id"]],
                            void_reason="superseded by reroll" if is_reroll else "superseded by revision",
                            voided_by=str(invocation.author.id),
                        )
                    # Delete old Discord messages
                    if old_ruling_channel and old_ruling_msg_ids:
                        for mid in old_ruling_msg_ids:
                            try:
                                old_msg = await old_ruling_channel.fetch_message(int(mid))
                                await old_msg.delete()
                            except (discord.NotFound, discord.Forbidden, ValueError):
                                pass

                # effective_channel: where the ruling should appear
                # Falls back to old_ruling_channel (from metadata) if output_channel
                # wasn't passed explicitly (e.g. revision invoked without config channel)
                effective_channel: discord.TextChannel = (
                    output_channel or old_ruling_channel or invocation.channel
                )
                in_dedicated = effective_channel.id != invocation.channel.id

                # --- Send ruling and capture Discord message IDs for future revisions ---
                ruling_label = (
                    " (Revised)" if revision_request else
                    " (Rerolled)" if is_reroll else ""
                )
                sent_msg_ids: List[str] = []
                body_msg_ids: List[str] = []  # ruling text messages (excludes header)
                if in_dedicated:
                    header = (
                        f"**{ruling_id}{ruling_label}**"
                        f" — {target_msg.jump_url}\n"
                        f"Action by **{target_msg.author.display_name}** · "
                        f"Adjudicated by {invocation.author.mention}"
                    )
                    m = await effective_channel.send(header)
                    sent_msg_ids.append(str(m.id))
                    for i in range(0, len(ruling), 1900):
                        m = await effective_channel.send(ruling[i:i + 1900])
                        sent_msg_ids.append(str(m.id))
                        body_msg_ids.append(str(m.id))
                    if not is_reroll:
                        label = "revised" if revision_request else "posted"
                        await invocation.reply(
                            f"✅ Ruling {label} in {effective_channel.mention}."
                        )
                elif is_replacement:
                    # Inline revision/reroll: send standalone so it replaces the original
                    for i in range(0, len(ruling), 1900):
                        m = await effective_channel.send(ruling[i:i + 1900])
                        sent_msg_ids.append(str(m.id))
                        body_msg_ids.append(str(m.id))
                    if revision_request:
                        await invocation.reply("✅ Ruling revised.")
                else:
                    # First-time inline: reply to invocation
                    for i in range(0, len(ruling), 1900):
                        if i == 0:
                            m = await invocation.reply(ruling[i:i + 1900])
                        else:
                            m = await invocation.channel.send(ruling[i:i + 1900])
                        sent_msg_ids.append(str(m.id))
                        body_msg_ids.append(str(m.id))

                # Add 🎲 reaction to first body message so the action author can reroll
                if body_msg_ids:
                    try:
                        first_body = await effective_channel.fetch_message(
                            int(body_msg_ids[0])
                        )
                        await first_body.add_reaction("🎲")
                    except (discord.NotFound, discord.Forbidden, discord.HTTPException):
                        pass

                # Populate in-memory cache so on_raw_reaction_add can handle rerolls
                cache_entry = {
                    "target_channel_id": target_msg.channel.id,
                    "target_msg_id": target_msg.id,
                    "action_author_id": target_msg.author.id,
                    "invocation_channel_id": invocation.channel.id,
                    "invocation_msg_id": invocation.id,
                    "focus_question": focus_question,
                }
                for mid in body_msg_ids:
                    self._ruling_msg_cache[mid] = cache_entry

                # --- Cache the ruling in the archive ---
                try:
                    print(f"{_C_SEND}[GM STORE →] embedding ruling ({len(ruling)} chars){_C_RST}", flush=True)
                    ruling_embed, ruling_embed_usage = await self.llm.embed(ruling)
                    self.state.add_usage("embedding", ruling_embed_usage)
                    print(f"{_C_RECV}[GM STORE ←] dims={len(ruling_embed)}{_C_RST}", flush=True)
                    entry_type = (
                        "ruling_revision" if revision_request else
                        "ruling_reroll" if is_reroll else
                        "ruling"
                    )
                    self.memory.add(
                        text=ruling,
                        embedding=ruling_embed,
                        metadata={
                            "source_message_id": str(invocation.id),
                            "author_id": str(invocation.author.id),
                            "author_name": invocation.author.display_name,
                            "author_discord_name": invocation.author.name,
                            "channel": invocation.channel.name,
                            "channel_id": str(invocation.channel.id),
                            "timestamp": invocation.created_at.isoformat(),
                            "year": int(self.state.current_year),
                            "entry_type": entry_type,
                            "ruling_id": ruling_id or "",
                            "ruled_on_message_id": str(target_msg.id),
                            "ruled_on_author_id": str(target_msg.author.id),
                            "ruled_on_author_name": target_msg.author.display_name,
                            "ruled_on_author_discord_name": target_msg.author.name,
                            "has_image": False,
                            "full_message_text": ruling,
                            "original_excerpt": ruling[:500],
                            "ruling_discord_message_ids": ",".join(sent_msg_ids),
                            "ruling_discord_channel_id": str(effective_channel.id),
                        },
                    )
                except Exception:
                    log.exception("Failed to cache GM ruling in archive")

        except Exception as e:
            log.exception("GM ruling failed")
            await invocation.reply(f"⚠️ GM ruling failed: `{e}`")
        finally:
            self._gm_inflight.discard(gm_key)

    # ------------------------------------------------------------------
    # War GM mode
    # ------------------------------------------------------------------
    def _war_help_text(self) -> str:
        p = self.prefix
        return "\n".join([
            "⚔️ **War GM mode** — run these inside the war's thread:",
            f"`{p} war start <title>` — begin a war in this thread *(GM)*",
            f"`{p} war side <Side> <TAG|@player> …` — register a side *(GM)*",
            f"`{p} war npc <TAG> [on|off]` — toggle AI control of a nation *(GM)*",
            f"`{p} war move <text|message-link>` — submit & adjudicate a move *(belligerents)*",
            f"`{p} war npcact <TAG> [guidance]` — make an AI nation act now *(GM)*",
            f"`{p} war status` — show sides + the current war-state digest",
            f"`{p} war end` — draft the war chronicle *(GM)*",
            f"`{p} war commit` — write the chronicle to the archive *(GM)*",
            f"`{p} war cancel` — abandon the war, write nothing *(GM)*",
            "Reply to a ruling to revise it *(GM)*, or 🎲-react it to reroll.",
        ])

    def _war_briefing(self, war: dict) -> str:
        """Compact human-readable briefing of the war's sides and belligerents."""
        lines = [
            f"War ID: {war['war_id']}",
            f"Title: {war.get('title', '')}",
            f"Current in-game year: {self.state.current_year} "
            f"(war began {war.get('start_year')})",
            "Sides & belligerents:",
        ]
        sides = war.get("sides", {})
        bel = war.get("belligerents", {})
        if sides:
            for side_name, side in sides.items():
                tag_strs = []
                for tag in side.get("tags", []):
                    npc = bel.get(tag, {}).get("npc", False)
                    tag_strs.append(f"{tag}{' [AI]' if npc else ''}")
                lines.append(f"  • {side_name}: {', '.join(tag_strs) or '(none)'}")
        else:
            lines.append("  (no sides registered yet)")
        return "\n".join(lines)

    def _render_war_log(self, war: dict, budget: int) -> str:
        """Render the verbatim war log (skipping superseded rulings), oldest→newest.

        If the rendered text exceeds `budget` chars, the oldest blocks are
        dropped (their effects survive in the maintained state digest).
        """
        blocks: list[str] = []
        for e in war.get("log", []):
            if e.get("superseded"):
                continue
            t = e.get("type")
            ts = str(e.get("timestamp", ""))[:19]
            if t in ("move", "npc_move"):
                who = e.get("tag", "?")
                side = e.get("side", "")
                actor = e.get("author_name") or who
                kind = "AI MOVE" if t == "npc_move" else "MOVE"
                head = (f"[{kind} #{e.get('seq')}] {who}"
                        f"{f' ({side})' if side else ''} — by {actor} @ {ts}")
                blocks.append(f"{head}\n{e.get('text', '')}")
            elif t == "ruling":
                blocks.append(
                    f"[RULING on move #{e.get('of_seq', '?')}]\n{e.get('text', '')}"
                )
        if not blocks:
            return "(No moves recorded yet — this is the opening move of the war.)"
        text = "\n\n".join(blocks)
        if len(text) <= budget:
            return text
        kept: list[str] = []
        total = 0
        for b in reversed(blocks):
            if kept and total + len(b) > budget:
                break
            kept.append(b)
            total += len(b)
        kept.reverse()
        note = ("[... earlier moves omitted from verbatim replay; their effects "
                "are captured in the STATE DIGEST above ...]")
        return note + "\n\n" + "\n\n".join(kept)

    async def _pull_nation_lore(self, tag: str, hint_text: str,
                                top_k: int, item_chars: int) -> str:
        """Pull a nation's established lore from the main archive for grounding."""
        query = (f"{tag} nation military forces capabilities army navy doctrine "
                 f"history leadership. {hint_text}")
        try:
            q_embed, usage = await self.llm.embed(
                _trim_context_text(query, _EMBED_INPUT_CHARS)
            )
            self.state.add_usage("embedding", usage)
            hits = self.memory.search(q_embed, top_k=top_k)
        except Exception:
            log.exception("War nation-lore retrieval failed for %s", tag)
            return ""
        if not hits:
            return ""
        lines = []
        for h in hits:
            meta = h["metadata"]
            full = meta.get("full_message_text", "") or meta.get("original_excerpt", "")
            full = _trim_context_text(full, item_chars)
            lines.append(
                f"[{str(meta.get('timestamp', ''))[:10]}, "
                f"by {meta.get('author_name', '?')}, #{meta.get('channel', '?')}, "
                f"year {meta.get('year', '?')}]\n"
                f"Summary: {h['text']}\n"
                f"Full: {full}"
            )
        return "\n\n".join(lines)

    # ---- Command dispatcher -----------------------------------------
    async def _cmd_war(self, message: discord.Message, arg: str):
        if not isinstance(message.channel, discord.Thread):
            await message.reply(
                f"⚔️ War mode runs inside a thread. Open a thread for the war, then "
                f"run `{self.prefix} war start <title>` there. (`{self.prefix} war help` "
                f"for the full command list.)"
            )
            return

        parts = arg.strip().split(maxsplit=1)
        sub = parts[0].lower() if parts and parts[0] else ""
        rest = parts[1].strip() if len(parts) > 1 else ""
        thread = message.channel
        tid = str(thread.id)

        if sub in ("", "help"):
            await message.reply(self._war_help_text()); return
        if sub == "start":
            await self._war_start(message, thread, tid, rest); return
        if sub == "status":
            await self._war_status(message, tid); return
        if sub == "move":
            await self._war_move(message, thread, tid, rest); return

        # Everything below is GM-only.
        if not self._is_war_gm(message.author):
            await message.reply("⚠️ Only a war GM can use that war command.")
            return
        if sub == "side":
            await self._war_side(message, thread, tid, rest)
        elif sub == "npc":
            await self._war_npc_toggle(message, tid, rest)
        elif sub == "npcact":
            await self._war_npcact(message, thread, tid, rest)
        elif sub == "end":
            await self._war_end(message, thread, tid)
        elif sub == "commit":
            await self._war_commit(message, thread, tid)
        elif sub == "cancel":
            await self._war_cancel(message, tid)
        else:
            await message.reply(
                f"Unknown war subcommand `{sub}`. Try `{self.prefix} war help`."
            )

    async def _war_start(self, message, thread, tid, title):
        if not self._is_war_gm(message.author):
            await message.reply("⚠️ Only a war GM can start a war."); return
        existing = self.war_store.get_active_war(tid)
        if existing:
            await message.reply(
                f"⚠️ A war is already live in this thread: **{existing['war_id']}** "
                f"({existing.get('title', '')}). End or cancel it first."
            )
            return
        year = self.state.current_year
        seq = self.state.next_war_id(year)
        war_id = f"WAR-{year}-{seq:04d}"
        title = (title or "").strip() or thread.name or war_id
        self.war_store.create_war(
            tid, war_id, title, str(message.guild.id), str(message.author.id), year
        )
        self.state.bump("wars_started")
        self.flog.log_command("war start", str(message.author), thread.name,
                              f"{war_id} {title}")
        await message.reply(
            f"⚔️ **{war_id}** — *{title}* — is now active in this thread, with "
            f"{message.author.mention} as GM.\n\n"
            f"**Next:** register the sides, e.g.\n"
            f"`{self.prefix} war side Allies USA GBR`\n"
            f"`{self.prefix} war side Coalition @player SYR`\n\n"
            f"Unoccupied nations are auto-flagged AI-controlled (toggle with "
            f"`{self.prefix} war npc <TAG> on|off`). Then players submit moves with "
            f"`{self.prefix} war move <description or message-link>`."
        )

    async def _war_side(self, message, thread, tid, rest):
        war = self.war_store.get_active_war(tid)
        if not war:
            await message.reply(
                f"⚠️ No active war here. Run `{self.prefix} war start` first."); return
        bits = rest.split()
        if len(bits) < 2:
            await message.reply(
                f"Usage: `{self.prefix} war side <SideName> <TAG|@player> [more…]`\n"
                f"Example: `{self.prefix} war side Allies USA GBR @player`"
            )
            return
        side_name = bits[0]
        occupation = self.map_store.get_occupation()
        belligerents: list[tuple[str, bool]] = []
        notes: list[str] = []
        for tok in bits[1:]:
            m = MENTION_RE.match(tok)
            if m:
                member = message.guild.get_member(int(m.group(1)))
                tag = parse_tag(member.display_name) if member else None
                if not tag:
                    notes.append(f"⚠️ couldn't read a nation TAG from {tok}")
                    continue
                belligerents.append((tag, False))  # mentioned member = human-controlled
            else:
                tag = tok.upper().strip(",")
                if not re.match(r"^[A-Z]{2,4}$", tag):
                    notes.append(f"⚠️ `{tok}` is not a TAG or @mention")
                    continue
                npc = tag not in occupation  # auto-flag unoccupied nations as AI
                belligerents.append((tag, npc))
        if not belligerents:
            await message.reply("⚠️ No valid belligerents parsed.\n" + "\n".join(notes))
            return
        self.war_store.add_side(tid, side_name, belligerents)
        desc = ", ".join(f"{t}{' [AI]' if npc else ''}" for t, npc in belligerents)
        suffix = ("\n" + "\n".join(notes)) if notes else ""
        await message.reply(
            f"✅ Side **{side_name}**: {desc}{suffix}\n"
            f"(Toggle AI control with `{self.prefix} war npc <TAG> on|off`.)"
        )

    async def _war_npc_toggle(self, message, tid, rest):
        war = self.war_store.get_active_war(tid)
        if not war:
            await message.reply("⚠️ No active war here."); return
        bits = rest.split()
        bel = war.get("belligerents", {})
        if not bits:
            if not bel:
                await message.reply("No belligerents registered yet."); return
            lines = ["**Belligerents:**"]
            for tag, info in bel.items():
                lines.append(
                    f"• `{tag}` — {info.get('side', '?')} — "
                    f"{'🤖 AI-controlled' if info.get('npc') else '🧑 player-controlled'}"
                )
            await message.reply("\n".join(lines)); return
        tag = bits[0].upper()
        if tag not in bel:
            await message.reply(f"⚠️ `{tag}` is not a belligerent in this war."); return
        if len(bits) >= 2:
            val = bits[1].lower() in ("on", "true", "yes", "1", "ai")
        else:
            val = not bel[tag].get("npc", False)
        self.war_store.set_npc(tid, tag, val)
        await message.reply(
            f"✅ `{tag}` is now {'🤖 AI-controlled' if val else '🧑 player-controlled'}."
        )

    async def _war_status(self, message, tid):
        war = self.war_store.get_war(tid)
        if not war:
            await message.reply("No war has been started in this thread."); return
        digest = war.get("state_digest") or "(no moves adjudicated yet)"
        body = (
            f"⚔️ **{war['war_id']}** — *{war.get('title', '')}* — status: "
            f"**{war.get('status')}**\n"
            f"{self._war_briefing(war)}\n\n"
            f"**State digest:**\n{digest}"
        )
        await self._send_chunked_reply(message, body)

    async def _resolve_war_move_text(self, message, rest):
        """Resolve move text from inline text, a message link, or a reply."""
        rest = (rest or "").strip()
        link = MSG_LINK_RE.search(rest)
        if link:
            try:
                ch = self.get_channel(int(link.group(2)))
                if ch is None:
                    ch = await self.fetch_channel(int(link.group(2)))
                tmsg = await ch.fetch_message(int(link.group(3)))
                return (await self._message_text_with_documents(tmsg)).strip(), tmsg
            except Exception:
                return "", None
        if rest:
            return rest, None
        if message.reference and message.reference.message_id:
            try:
                tmsg = await message.channel.fetch_message(message.reference.message_id)
                return (await self._message_text_with_documents(tmsg)).strip(), tmsg
            except Exception:
                return "", None
        return "", None

    async def _war_move(self, message, thread, tid, rest):
        war = self.war_store.get_active_war(tid)
        if not war:
            await message.reply(
                f"⚠️ No active war in this thread. A GM must run "
                f"`{self.prefix} war start` first."); return
        if war.get("status") == "pending_commit":
            await message.reply(
                "⚠️ This war has ended and is awaiting its chronicle commit — "
                "no further moves."); return
        move_text, src_msg = await self._resolve_war_move_text(message, rest)
        if not move_text:
            await message.reply(
                f"Usage: `{self.prefix} war move <what your nation does>` — "
                f"or paste a message link, or reply to a post with "
                f"`{self.prefix} war move`."); return

        bel = war.get("belligerents", {})
        actor_tag = parse_tag(message.author.display_name)
        if actor_tag not in bel:
            if self._is_war_gm(message.author):
                await message.reply(
                    f"⚠️ Your nickname's TAG isn't a registered belligerent. As GM, "
                    f"use `{self.prefix} war npcact <TAG>` to act for an AI nation, "
                    f"or register the side with `{self.prefix} war side`."
                )
            else:
                await message.reply(
                    "⚠️ You don't control a nation in this war (no registered TAG "
                    "matches your nickname). Ask the GM to add your side."
                )
            return
        actor_side = bel[actor_tag]["side"]
        await self._produce_war_ruling(
            thread=thread, tid=tid, invocation=message,
            actor_tag=actor_tag, actor_side=actor_side, is_npc=False,
            move_text=move_text,
            source_message_id=str((src_msg or message).id),
            author_id=str(message.author.id),
            author_name=message.author.display_name,
            allow_npc_response=True,
        )

    async def _war_npcact(self, message, thread, tid, rest):
        war = self.war_store.get_active_war(tid)
        if not war:
            await message.reply("⚠️ No active war here."); return
        bits = rest.split(maxsplit=1)
        if not bits:
            await message.reply(
                f"Usage: `{self.prefix} war npcact <TAG> [guidance]`"); return
        tag = bits[0].upper()
        guidance = bits[1].strip() if len(bits) > 1 else None
        bel = war.get("belligerents", {})
        if tag not in bel:
            await message.reply(
                f"⚠️ `{tag}` is not a belligerent. Add it with "
                f"`{self.prefix} war side`."); return
        if not bel[tag].get("npc"):
            await message.reply(
                f"⚠️ `{tag}` is player-controlled. Run "
                f"`{self.prefix} war npc {tag} on` first if you want the bot to act "
                f"for it."); return
        await self._produce_war_npc_move(thread, tid, tag,
                                         trigger="a GM prompt", guidance=guidance)

    async def _war_cancel(self, message, tid):
        war = self.war_store.get_active_war(tid)
        if not war:
            await message.reply("⚠️ No active war in this thread."); return
        self.war_store.set_status(tid, "cancelled")
        self.flog.log_command("war cancel", str(message.author),
                              getattr(message.channel, "name", "?"), war["war_id"])
        await message.reply(
            f"🚫 War **{war['war_id']}** cancelled. Its move log is retained for "
            f"records, but nothing was written to the archive."
        )

    # ---- Adjudication engine ----------------------------------------
    async def _produce_war_ruling(self, *, thread, tid, invocation,
                                  actor_tag, actor_side, is_npc, move_text,
                                  source_message_id, author_id, author_name,
                                  allow_npc_response, guidance=None,
                                  revision_request=None, existing_move_seq=None,
                                  old_ruling_seq=None, old_ruling_msg_ids=None,
                                  is_reroll=False):
        war_key = (f"{tid}:{existing_move_seq or source_message_id}:"
                   f"{'rev' if revision_request else 'rr' if is_reroll else 'new'}")
        if war_key in self._war_inflight:
            log.info("Skipping duplicate war ruling %s", war_key)
            return
        self._war_inflight.add(war_key)
        try:
            async with thread.typing():
                war = self.war_store.get_active_war(tid)
                if not war:
                    return
                ctx = self._war_cfg.get("context", {})
                budget = int(ctx.get("log_char_budget", 120000))
                arch_top_k = int(ctx.get("archive_top_k", 12))
                arch_item = int(ctx.get("archive_item_chars", 4000))

                # 1. Record the move (unless this is a revision/reroll of one)
                if existing_move_seq is None:
                    move_seq = self.war_store.append_log(tid, {
                        "type": "npc_move" if is_npc else "move",
                        "tag": actor_tag, "side": actor_side, "npc": is_npc,
                        "author_id": author_id, "author_name": author_name,
                        "text": move_text, "source_message_id": source_message_id,
                    })
                    war = self.war_store.get_active_war(tid) or war
                else:
                    move_seq = existing_move_seq

                # 2. Acting nation's established lore from the main archive
                nation_lore = await self._pull_nation_lore(
                    actor_tag, move_text, arch_top_k, arch_item
                )

                # 3. Build context: briefing + digest + verbatim log + lore + move
                digest = war.get("state_digest") or (
                    "(No state digest yet — this is early in the war.)")
                log_text = self._render_war_log(war, budget)
                parts = [
                    "=== WAR BRIEFING ===", self._war_briefing(war), "",
                    "=== WAR STATE DIGEST ===", digest, "",
                    "=== WAR LOG (verbatim, oldest → newest) ===", log_text, "",
                    f"=== ACTING NATION: {actor_tag} ({actor_side})"
                    f"{' — AI-CONTROLLED' if is_npc else ''} ===",
                    f"--- {actor_tag} established lore from the campaign archive ---",
                    nation_lore or "(No archive lore found for this nation.)", "",
                    f"=== MOVE TO ADJUDICATE (this is move #{move_seq}) ===",
                    move_text,
                ]
                if guidance:
                    parts += ["", "=== GM GUIDANCE ===", guidance]
                if revision_request:
                    parts += [
                        "", "=== GM REVISION REQUEST ===",
                        "A GM has asked you to revise your prior ruling on this move. "
                        "Re-issue the ruling in full, adjusted per their direction:",
                        revision_request,
                    ]
                user_content = "\n".join(parts)

                msgs = [
                    {"role": "system", "content": self.prompts.get("war_move_ruling")},
                    {"role": "user", "content": user_content},
                ]
                _ws = self._mode_settings("war")
                print(f"{_C_SEND}[WAR      →] move#{move_seq} {actor_tag}"
                      f"{' [AI]' if is_npc else ''} model={_ws['model'] or 'default'}"
                      f"  thinking={_fmt_think(_ws['thinking_budget'])}{_C_RST}",
                      flush=True)
                ruling, usage = await self.llm.chat_messages(msgs, **_ws)
                self.state.add_usage("war", usage)

                # 4. Supersede the prior ruling (revision/reroll) and delete its posts
                if old_ruling_seq is not None:
                    self.war_store.mark_superseded(tid, old_ruling_seq)
                if old_ruling_msg_ids:
                    for mid in old_ruling_msg_ids:
                        try:
                            om = await thread.fetch_message(int(mid))
                            await om.delete()
                        except (discord.NotFound, discord.Forbidden,
                                discord.HTTPException, ValueError):
                            pass

                # 5. Post the ruling
                label = (" (Revised)" if revision_request else
                         " (Rerolled)" if is_reroll else "")
                actor_label = f"AI · {actor_tag}" if is_npc else actor_tag
                header = (f"⚔️ **{war['war_id']} — Ruling on move #{move_seq}** "
                          f"({actor_label}){label}")
                sent_ids: list[str] = []
                body_ids: list[str] = []
                m = await thread.send(header)
                sent_ids.append(str(m.id))
                for i in range(0, len(ruling), 1900):
                    mm = await thread.send(ruling[i:i + 1900])
                    sent_ids.append(str(mm.id))
                    body_ids.append(str(mm.id))

                # 6. Record the ruling in the war log
                ruling_seq = self.war_store.append_log(tid, {
                    "type": "ruling", "of_seq": move_seq, "text": ruling,
                    "discord_message_ids": sent_ids, "revision_of": old_ruling_seq,
                })

                if is_npc:
                    self.state.bump("war_npc_moves")
                if not (revision_request or is_reroll):
                    self.state.bump("war_moves_adjudicated")

                # 7. 🎲 reroll reaction + re-adjudication cache
                if body_ids:
                    try:
                        fb = await thread.fetch_message(int(body_ids[0]))
                        await fb.add_reaction("🎲")
                    except (discord.NotFound, discord.Forbidden,
                            discord.HTTPException):
                        pass
                cache_entry = {
                    "thread_id": tid, "war_id": war["war_id"],
                    "move_seq": move_seq, "ruling_seq": ruling_seq,
                    "move_text": move_text, "actor_tag": actor_tag,
                    "actor_side": actor_side, "is_npc": is_npc,
                    "author_id": author_id, "author_name": author_name,
                    "source_message_id": source_message_id, "guidance": guidance,
                    "ruling_msg_ids": sent_ids,
                }
                for mid in sent_ids:
                    self._war_ruling_cache[mid] = cache_entry

                # 8. Update the running state digest incrementally
                await self._update_war_digest(
                    tid, war, move_seq, actor_tag, is_npc, move_text, ruling, ruling_seq
                )

            # 9. Auto-generate AI responses for targeted NPC nations
            if (allow_npc_response and not is_npc
                    and self._war_cfg.get("npc_auto_respond", True)):
                await self._maybe_npc_responses(
                    thread, tid, actor_tag, actor_side, move_text, ruling
                )
        except Exception as e:
            log.exception("War ruling failed")
            if invocation is not None:
                try:
                    await invocation.reply(f"⚠️ War ruling failed: `{e}`")
                except discord.HTTPException:
                    pass
        finally:
            self._war_inflight.discard(war_key)

    async def _update_war_digest(self, tid, war, move_seq, actor_tag, is_npc,
                                 move_text, ruling, ruling_seq):
        """Regenerate the war-state digest after a ruling (small/cheap model)."""
        try:
            current = war.get("state_digest") or "(no digest yet)"
            user = (
                f"Current digest:\n{current}\n\n"
                f"NEW MOVE #{move_seq} by {actor_tag}{' [AI]' if is_npc else ''}:\n"
                f"{move_text}\n\n"
                f"RULING:\n{ruling}\n\n"
                f"Produce the updated digest."
            )
            _ds = self._mode_settings("war_digest")
            digest, usage = await self.llm.chat(
                self.prompts.get("war_state_digest"),
                _trim_context_text(user, _GM_WIDER_CONTEXT_CHARS), **_ds,
            )
            self.state.add_usage("war", usage)
            if digest.strip():
                self.war_store.set_digest(tid, digest.strip(), ruling_seq)
        except Exception:
            log.exception("War digest update failed for %s", tid)

    async def _maybe_npc_responses(self, thread, tid, actor_tag, actor_side,
                                   move_text, ruling):
        """Generate moves for AI nations on other sides that this move targets."""
        war = self.war_store.get_active_war(tid)
        if not war:
            return
        bel = war.get("belligerents", {})
        cap = int(self._war_cfg.get("npc_max_responses_per_move", 2))
        haystack = f"{move_text}\n{ruling}".upper()
        targets = []
        for tag, info in bel.items():
            if not info.get("npc"):
                continue
            if info.get("side") == actor_side:
                continue
            if re.search(rf"\b{re.escape(tag)}\b", haystack):
                targets.append(tag)
        for tag in targets[:cap]:
            await self._produce_war_npc_move(
                thread, tid, tag,
                trigger=f"{actor_tag}'s move and its outcome", guidance=None,
            )

    async def _produce_war_npc_move(self, thread, tid, npc_tag,
                                    trigger=None, guidance=None):
        """Have the bot declare an AI nation's war move, then adjudicate it."""
        war = self.war_store.get_active_war(tid)
        if not war:
            return
        bel = war.get("belligerents", {})
        if npc_tag not in bel:
            return
        side = bel[npc_tag]["side"]
        ctx = self._war_cfg.get("context", {})
        budget = int(ctx.get("log_char_budget", 120000))
        arch_top_k = int(ctx.get("archive_top_k", 12))
        arch_item = int(ctx.get("archive_item_chars", 4000))
        try:
            async with thread.typing():
                nation_lore = await self._pull_nation_lore(
                    npc_tag, war.get("title", ""), arch_top_k, arch_item
                )
                digest = war.get("state_digest") or "(early in the war)"
                parts = [
                    "=== WAR BRIEFING ===", self._war_briefing(war), "",
                    "=== WAR STATE DIGEST ===", digest, "",
                    "=== WAR LOG (verbatim, oldest → newest) ===",
                    self._render_war_log(war, budget), "",
                    f"=== YOU COMMAND: {npc_tag} ({side}) — AI-CONTROLLED ===",
                    f"--- {npc_tag} established lore from the campaign archive ---",
                    nation_lore or "(No archive lore found for this nation.)", "",
                ]
                if trigger:
                    parts += [f"You are reacting to: {trigger}.", ""]
                if guidance:
                    parts += ["=== GM GUIDANCE ===", guidance, ""]
                parts += ["Declare this nation's next war move now."]
                _ns = self._mode_settings("war_npc")
                print(f"{_C_SEND}[WAR NPC  →] {npc_tag} ({side}) "
                      f"model={_ns['model'] or 'default'}{_C_RST}", flush=True)
                move_text, usage = await self.llm.chat(
                    self.prompts.get("war_npc_move"), "\n".join(parts), **_ns,
                )
                self.state.add_usage("war", usage)
                move_text = (move_text or "").strip()
                if not move_text:
                    return
                await thread.send(f"🤖 **{npc_tag}** ({side}) responds:")
                for i in range(0, len(move_text), 1900):
                    await thread.send(move_text[i:i + 1900])
        except Exception:
            log.exception("War NPC move generation failed for %s", npc_tag)
            return

        # Adjudicate the NPC's declared move (no further NPC cascade)
        synthetic_id = f"npc-{tid}-{npc_tag}-{datetime.now(timezone.utc).timestamp()}"
        await self._produce_war_ruling(
            thread=thread, tid=tid, invocation=None,
            actor_tag=npc_tag, actor_side=side, is_npc=True,
            move_text=move_text, source_message_id=synthetic_id,
            author_id="", author_name=f"AI/{npc_tag}",
            allow_npc_response=False,
        )

    # ---- War end / chronicle ----------------------------------------
    def _split_chronicle(self, text: str) -> list[dict]:
        """Split chronicle text on '=== ENTRY: <title> ===' markers."""
        marker = re.compile(r"(?m)^===\s*ENTRY:\s*(.+?)\s*===\s*$")
        matches = list(marker.finditer(text))
        if not matches:
            return [{"title": "War Chronicle", "text": text.strip()}]
        entries = []
        for i, mt in enumerate(matches):
            start = mt.end()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            body = text[start:end].strip()
            if body:
                entries.append({"title": mt.group(1).strip(), "text": body})
        return entries or [{"title": "War Chronicle", "text": text.strip()}]

    async def _war_end(self, message, thread, tid):
        war = self.war_store.get_active_war(tid)
        if not war:
            await message.reply("⚠️ No active war in this thread."); return
        if not any(e.get("type") == "ruling" for e in war.get("log", [])):
            await message.reply(
                "⚠️ This war has no adjudicated moves yet — nothing to chronicle."); return
        async with thread.typing():
            user = "\n".join([
                "=== WAR BRIEFING ===", self._war_briefing(war), "",
                "=== FINAL STATE DIGEST ===",
                war.get("state_digest") or "(no digest)", "",
                "=== COMPLETE WAR LOG (verbatim, oldest → newest) ===",
                self._render_war_log(war, 1_000_000),
            ])
            _cs = self._mode_settings("war_chronicle")
            print(f"{_C_SEND}[WAR END  →] chronicle for {war['war_id']} "
                  f"model={_cs['model'] or 'default'}{_C_RST}", flush=True)
            chronicle, usage = await self.llm.chat(
                self.prompts.get("war_chronicle"),
                _trim_context_text(user, _GM_PRIOR_CONTEXT_CHARS), **_cs,
            )
            self.state.add_usage("war", usage)
        entries = self._split_chronicle(chronicle)
        self.war_store.set_chronicle_draft(tid, entries)
        self.war_store.set_status(tid, "pending_commit")
        self.flog.log_command("war end", str(message.author), thread.name,
                              war["war_id"])
        plural = "entry" if len(entries) == 1 else "entries"
        await thread.send(
            f"📜 **Draft chronicle for {war['war_id']} — {war.get('title', '')}** "
            f"({len(entries)} {plural}). Review below, then a GM runs "
            f"`{self.prefix} war commit` to write it to the archive — or "
            f"`{self.prefix} war cancel` to discard. (Re-run `{self.prefix} war end` "
            f"to regenerate the draft.)"
        )
        for e in entries:
            await thread.send(f"**{e['title']}**")
            for i in range(0, len(e["text"]), 1900):
                await thread.send(e["text"][i:i + 1900])

    async def _war_commit(self, message, thread, tid):
        war = self.war_store.get_war(tid)
        if not war:
            await message.reply("⚠️ No war in this thread."); return
        if war.get("status") != "pending_commit":
            await message.reply(
                f"⚠️ This war isn't awaiting commit (status: {war.get('status')}). "
                f"Run `{self.prefix} war end` first."); return
        entries = war.get("chronicle_draft", [])
        if not entries:
            await message.reply(
                f"⚠️ No chronicle draft found. Run `{self.prefix} war end` again."); return
        year = self.state.current_year
        participants = ",".join(sorted(war.get("belligerents", {}).keys()))
        stored = 0
        async with thread.typing():
            for idx, e in enumerate(entries, 1):
                text = f"{e['title']}\n\n{e['text']}"
                try:
                    embed, usage = await self.llm.embed(
                        _trim_context_text(text, _EMBED_INPUT_CHARS)
                    )
                    self.state.add_usage("embedding", usage)
                    self.memory.add(
                        text=text, embedding=embed,
                        metadata={
                            "source_message_id": f"{war['war_id']}-chronicle-{idx}",
                            "author_id": "", "author_name": "War Chronicle",
                            "author_discord_name": "DNC",
                            "channel": thread.name, "channel_id": tid,
                            "timestamp": datetime.now(timezone.utc).isoformat(),
                            "year": int(year), "entry_type": "war_chronicle",
                            "war_id": war["war_id"], "war_title": war.get("title", ""),
                            "war_participants": participants, "war_thread_id": tid,
                            "chronicle_part": idx, "chronicle_total": len(entries),
                            "has_image": False,
                            "full_message_text": text, "original_excerpt": text[:500],
                        },
                    )
                    stored += 1
                except Exception:
                    log.exception("Failed to store war chronicle entry %d", idx)
        self.war_store.set_status(tid, "ended")
        self.state.bump("wars_ended")
        self.flog.log_command("war commit", str(message.author), thread.name,
                              f"{war['war_id']} {stored} entries")
        plural = "entry" if stored == 1 else "entries"
        await message.reply(
            f"✅ Committed **{stored}** chronicle {plural} for **{war['war_id']}** to "
            f"the archive (total memories now {self.memory.count()}). This war is "
            f"now closed."
        )

    # ------------------------------------------------------------------
    # Memory ingestion
    # ------------------------------------------------------------------
    async def _extract_doc_text(self, attachment: discord.Attachment) -> str:
        """Download and extract plain text from a PDF or DOCX attachment."""
        doc_type = _doc_attachment_type(attachment)
        if not doc_type:
            return ""
        data = await attachment.read()
        try:
            if doc_type == "pdf":
                import pypdf
                reader = pypdf.PdfReader(io.BytesIO(data))
                return "\n".join(page.extract_text() or "" for page in reader.pages)
            else:
                import docx as _docx
                doc = _docx.Document(io.BytesIO(data))
                blocks: list[str] = []
                blocks.extend(p.text for p in doc.paragraphs if p.text.strip())
                for table in doc.tables:
                    for row in table.rows:
                        cells = [
                            _collapse_blank_lines(cell.text)
                            for cell in row.cells
                            if cell.text.strip()
                        ]
                        if cells:
                            blocks.append(" | ".join(cells))
                return _collapse_blank_lines("\n".join(blocks))
        except Exception:
            log.exception("Failed to extract text from %s (%s)", attachment.filename, doc_type)
            return ""

    async def _extract_document_parts(
        self, attachments: list[discord.Attachment]
    ) -> list[tuple[str, str]]:
        parts: list[tuple[str, str]] = []
        for a in attachments:
            print(f"{_C_SEND}[DOC     →] extracting {a.filename!r}{_C_RST}", flush=True)
            text = await self._extract_doc_text(a)
            if text.strip():
                clean = _collapse_blank_lines(text)
                parts.append((a.filename or "document", clean))
                print(f"{_C_RECV}[DOC     ←] {len(clean)} chars from {a.filename!r}{_C_RST}", flush=True)
            else:
                print(f"{_C_ERR}[DOC     ✗] no text extracted from {a.filename!r}{_C_RST}", flush=True)
        return parts

    def _format_document_context(self, parts: list[tuple[str, str]]) -> str:
        if not parts:
            return ""
        return "\n\n[Attached Document Text]:\n" + "\n\n".join(
            f"[{filename}]:\n{text}" for filename, text in parts
        )

    async def _message_text_with_documents(self, message: discord.Message) -> str:
        text = (message.content or "").strip()
        doc_attachments = [a for a in message.attachments if _doc_attachment_type(a)]
        if not doc_attachments:
            return text
        doc_parts = await self._extract_document_parts(doc_attachments)
        doc_context = self._format_document_context(doc_parts)
        return f"{text}{doc_context}".strip()

    async def _store_document_chunks(
        self,
        *,
        doc_parts: list[tuple[str, str]],
        source_message_id: str,
        group_message_ids: str,
        author_id: str,
        author_name: str,
        author_discord_name: str,
        channel_name: str,
        channel_id: str,
        timestamp: str,
        year: int,
    ) -> int:
        stored = 0
        for filename, doc_text in doc_parts:
            chunks = list(_iter_text_chunks(doc_text))
            total = len(chunks)
            for idx, chunk in enumerate(chunks, start=1):
                chunk_text = (
                    f"Document chunk {idx}/{total} from {filename}.\n"
                    f"Author: {author_name}\n"
                    f"Channel: #{channel_name}\n"
                    f"In-game year: {year}\n\n"
                    f"{chunk}"
                )
                print(f"{_C_SEND}[DOC EMB →] {filename!r} chunk {idx}/{total} ({len(chunk)} chars){_C_RST}", flush=True)
                embedding, embed_usage = await self.llm.embed(chunk_text)
                self.state.add_usage("embedding", embed_usage)
                self.memory.add(
                    text=chunk_text,
                    embedding=embedding,
                    metadata={
                        "source_message_id": source_message_id,
                        "group_message_ids": group_message_ids,
                        "author_id": author_id,
                        "author_name": author_name,
                        "author_discord_name": author_discord_name,
                        "channel": channel_name,
                        "channel_id": channel_id,
                        "timestamp": timestamp,
                        "year": int(year),
                        "entry_type": "document_chunk",
                        "document_filename": filename,
                        "document_chunk_index": idx,
                        "document_chunk_total": total,
                        "has_image": False,
                        "full_message_text": chunk,
                        "original_excerpt": chunk[:500],
                    },
                )
                stored += 1
        return stored

    async def _ingest_message(self, message: discord.Message, year_override: int | None = None) -> bool:
        log.info("Ingesting #%s msg %s (%d chars)",
                 message.channel.name, message.id, len(message.content))
        print(f"\n{_C_HDR}{'─'*70}{_C_RST}", flush=True)
        print(f"{_C_HDR}[INGEST] msg={message.id}  author={message.author.display_name!r}"
              f"  #{message.channel.name}  len={len(message.content)}{_C_RST}", flush=True)

        # 1. Handle vision/OCR if images are attached
        image_context = ""
        images = [a for a in message.attachments if a.content_type and a.content_type.startswith("image/")]
        if images:
            try:
                vision_prompt = self.prompts.get("vision_extraction")
                urls = [a.url for a in images]
                log.info("Extracting vision context from %d images", len(urls))
                print(f"{_C_SEND}[VISION  →] model={self.llm.vision_model}  images={len(urls)}{_C_RST}", flush=True)
                for u in urls:
                    print(f"{_C_DIM}            {u}{_C_RST}", flush=True)
                image_context, usage = await self.llm.vision_extract(urls, vision_prompt)
                self.state.add_usage("vision", usage)
                preview = image_context.replace("\n", "↵")[:600]
                print(f"{_C_RECV}[VISION  ←] tokens=p{usage['prompt_tokens']}/c{usage['completion_tokens']}{_C_RST}", flush=True)
                print(f"{_C_DIM}            {preview}{_C_RST}", flush=True)
                image_context = f"\n\n[Attached Image Context/OCR]:\n{image_context}"
            except Exception:
                log.exception("Vision extraction failed for %s", message.id)
                print(f"{_C_ERR}[VISION  ✗] extraction failed — see log above{_C_RST}", flush=True)
        else:
            print(f"{_C_DIM}[VISION]    no images attached{_C_RST}", flush=True)

        # 2. Handle PDF / DOCX attachments
        doc_context = ""
        doc_attachments = [a for a in message.attachments if _doc_attachment_type(a)]
        doc_parts: list[tuple[str, str]] = []
        if doc_attachments:
            doc_parts = await self._extract_document_parts(doc_attachments)
            doc_context = self._format_document_context(doc_parts)

        char_name = message.author.display_name
        acct_name = message.author.name
        author_line = (
            f"Author: {char_name} (Discord: @{acct_name})"
            if char_name != acct_name
            else f"Author: {char_name}"
        )
        ingest_year = year_override if year_override is not None else self.state.current_year
        prompt_input = (
            f"{author_line}\n"
            f"Channel: #{message.channel.name}\n"
            f"In-game year: {ingest_year}\n"
            f"---\n{message.content}{image_context}{doc_context}"
        )
        archivist_input = _trim_context_text(prompt_input, _ARCHIVIST_INPUT_CHARS)
        _arch = self._mode_settings("archivist")
        print(f"{_C_SEND}[ARCHIVST→] model={_arch['model'] or 'default'}  temp={_arch['temperature']}"
              f"  thinking={_fmt_think(_arch['thinking_budget'])}  input={len(archivist_input)} chars{_C_RST}", flush=True)
        print(f"{_C_DIM}            {archivist_input[:500].replace(chr(10), '↵')}{_C_RST}", flush=True)

        self.state.bump("messages_sent_to_archivist")
        raw_archivist, chat_usage = await self.llm.chat(
            self.prompts.get("memory_extraction"), archivist_input,
            **self._mode_settings("archivist"),
        )
        self.state.add_usage("archivist", chat_usage)
        print(f"{_C_RECV}[ARCHIVST←] tokens=p{chat_usage['prompt_tokens']}/c{chat_usage['completion_tokens']}{_C_RST}", flush=True)
        print(f"{_C_DIM}            {raw_archivist}{_C_RST}", flush=True)

        if (
            self.cfg["memory"].get("filter_non_lore", True)
            and raw_archivist.strip().upper().startswith("NO_LORE")
        ):
            self.state.bump("messages_filtered_non_lore")
            self.flog.log_non_lore(
                str(message.id), message.author.display_name,
                message.channel.name, len(message.content),
            )
            print(f"{_C_DIM}[EMBED]     skipped — NO_LORE filter{_C_RST}", flush=True)
            print(f"{_C_HDR}{'─'*70}{_C_RST}", flush=True)
            return False

        summary, tags = _parse_archivist_output(raw_archivist)
        if tags.get("action_type") or tags.get("entities") or tags.get("claims"):
            print(f"{_C_DIM}            tags: type={tags.get('action_type', '')!r} "
                  f"entities={tags.get('entities', '')[:80]!r} "
                  f"claims={tags.get('claims', '')[:80]!r}{_C_RST}", flush=True)

        print(f"{_C_SEND}[EMBED    →] model={self.llm.embedding_model}  text={len(summary)} chars{_C_RST}", flush=True)
        embedding, embed_usage = await self.llm.embed(summary)
        self.state.add_usage("embedding", embed_usage)
        print(f"{_C_RECV}[EMBED    ←] dims={len(embedding)}  tokens={embed_usage['total_tokens']}{_C_RST}", flush=True)

        store_full = self.cfg["memory"].get("store_full_message", True)
        full_text = f"{message.content}{image_context}{doc_context}" if store_full else ""

        self.memory.add(
            text=summary,
            embedding=embedding,
            metadata={
                "source_message_id": str(message.id),
                "author_id": str(message.author.id),
                "author_name": message.author.display_name,
                "author_discord_name": message.author.name,
                "channel": message.channel.name,
                "channel_id": str(message.channel.id),
                "timestamp": message.created_at.isoformat(),
                "year": int(ingest_year),
                "entry_type": "message",
                "has_image": bool(message.attachments),
                "full_message_text": full_text,
                "original_excerpt": message.content[:500],
                "entities": tags.get("entities", ""),
                "action_type": tags.get("action_type", ""),
                "claims": tags.get("claims", ""),
            },
        )
        if doc_parts:
            doc_chunks = await self._store_document_chunks(
                doc_parts=doc_parts,
                source_message_id=str(message.id),
                group_message_ids=str(message.id),
                author_id=str(message.author.id),
                author_name=message.author.display_name,
                author_discord_name=message.author.name,
                channel_name=message.channel.name,
                channel_id=str(message.channel.id),
                timestamp=message.created_at.isoformat(),
                year=int(ingest_year),
            )
            if doc_chunks:
                print(f"{_C_OK}[DOC STORE✓] {doc_chunks} document chunks saved{_C_RST}", flush=True)
        self.state.bump("messages_archived")
        self.flog.log_archived(
            str(message.id), message.author.display_name,
            message.channel.name, len(message.content),
            int(ingest_year),
        )
        print(f"{_C_OK}[STORE    ✓] memory saved  total={self.memory.count()}{_C_RST}", flush=True)
        print(f"{_C_HDR}{'─'*70}{_C_RST}", flush=True)
        return True

    # ------------------------------------------------------------------
    # Post-chain grouping helpers
    # ------------------------------------------------------------------
    async def _queue_for_ingestion(self, message: discord.Message):
        """Buffer message; flush as a combined memory after chain_delay_seconds of silence."""
        key = (message.author.id, message.channel.id)

        existing = self._pending_timers.pop(key, None)
        if existing:
            existing.cancel()

        self._pending_chains.setdefault(key, []).append(message)
        self._pending_timers[key] = asyncio.create_task(self._delayed_flush(key))

    async def _delayed_flush(self, key: tuple[int, int]):
        await asyncio.sleep(self._chain_delay)
        await self._flush_chain(key)

    async def _flush_chain(self, key: tuple[int, int]):
        messages = self._pending_chains.pop(key, [])
        self._pending_timers.pop(key, None)
        if not messages:
            return
        try:
            if len(messages) == 1:
                await self._ingest_message(messages[0])
            else:
                await self._ingest_message_chain(messages)
        except Exception:
            log.exception("Memory ingestion failed for chain starting at %s", messages[0].id)

    async def _ingest_message_chain(self, messages: list[discord.Message],
                                     year_override: int | None = None) -> bool:
        """Ingest multiple consecutive messages from the same author as one combined memory."""
        primary = messages[0]
        combined_content = "\n\n".join(
            m.content.strip() for m in messages if (m.content or "").strip()
        )
        all_images = [
            a for m in messages for a in m.attachments
            if a.content_type and a.content_type.startswith("image/")
        ]
        all_docs = [
            a for m in messages for a in m.attachments
            if _doc_attachment_type(a)
        ]
        group_ids = [str(m.id) for m in messages]

        log.info("Ingesting chain of %d msgs from %s in #%s (%d chars)",
                 len(messages), primary.author.display_name, primary.channel.name,
                 len(combined_content))
        print(f"\n{_C_HDR}{'─'*70}{_C_RST}", flush=True)
        ids_preview = ",".join(group_ids[:3]) + ("…" if len(group_ids) > 3 else "")
        print(f"{_C_HDR}[CHAIN  ] msgs={len(messages)}  ids=[{ids_preview}]"
              f"  author={primary.author.display_name!r}"
              f"  #{primary.channel.name}  len={len(combined_content)}{_C_RST}", flush=True)

        image_context = ""
        if all_images:
            try:
                vision_prompt = self.prompts.get("vision_extraction")
                urls = [a.url for a in all_images]
                print(f"{_C_SEND}[VISION  →] model={self.llm.vision_model}  images={len(urls)}{_C_RST}", flush=True)
                image_context, usage = await self.llm.vision_extract(urls, vision_prompt)
                self.state.add_usage("vision", usage)
                image_context = f"\n\n[Attached Image Context/OCR]:\n{image_context}"
            except Exception:
                log.exception("Vision extraction failed for chain %s", group_ids[0])
                print(f"{_C_ERR}[VISION  ✗] extraction failed — see log above{_C_RST}", flush=True)
        else:
            print(f"{_C_DIM}[VISION]    no images attached{_C_RST}", flush=True)

        # 2. Handle PDF / DOCX attachments across the chain
        doc_context = ""
        doc_parts: list[tuple[str, str]] = []
        if all_docs:
            doc_parts = await self._extract_document_parts(all_docs)
            doc_context = self._format_document_context(doc_parts)

        char_name = primary.author.display_name
        acct_name = primary.author.name
        author_line = (
            f"Author: {char_name} (Discord: @{acct_name})"
            if char_name != acct_name else f"Author: {char_name}"
        )
        ingest_year = year_override if year_override is not None else self.state.current_year
        prompt_input = (
            f"{author_line}\n"
            f"Channel: #{primary.channel.name}\n"
            f"In-game year: {ingest_year}\n"
            f"---\n{combined_content}{image_context}{doc_context}"
        )
        archivist_input = _trim_context_text(prompt_input, _ARCHIVIST_INPUT_CHARS)
        _arch = self._mode_settings("archivist")
        print(f"{_C_SEND}[ARCHIVST→] model={_arch['model'] or 'default'}  temp={_arch['temperature']}"
              f"  thinking={_fmt_think(_arch['thinking_budget'])}  input={len(archivist_input)} chars{_C_RST}", flush=True)

        self.state.bump("messages_sent_to_archivist")
        raw_archivist, chat_usage = await self.llm.chat(
            self.prompts.get("memory_extraction"), archivist_input,
            **self._mode_settings("archivist"),
        )
        self.state.add_usage("archivist", chat_usage)
        print(f"{_C_RECV}[ARCHIVST←] tokens=p{chat_usage['prompt_tokens']}/c{chat_usage['completion_tokens']}{_C_RST}", flush=True)
        print(f"{_C_DIM}            {raw_archivist}{_C_RST}", flush=True)

        if (
            self.cfg["memory"].get("filter_non_lore", True)
            and raw_archivist.strip().upper().startswith("NO_LORE")
        ):
            self.state.bump("messages_filtered_non_lore")
            self.flog.log_non_lore(
                group_ids[0], primary.author.display_name,
                primary.channel.name, len(combined_content),
            )
            print(f"{_C_DIM}[EMBED]     skipped — NO_LORE filter{_C_RST}", flush=True)
            print(f"{_C_HDR}{'─'*70}{_C_RST}", flush=True)
            return False

        summary, tags = _parse_archivist_output(raw_archivist)
        if tags.get("action_type") or tags.get("entities") or tags.get("claims"):
            print(f"{_C_DIM}            tags: type={tags.get('action_type', '')!r} "
                  f"entities={tags.get('entities', '')[:80]!r} "
                  f"claims={tags.get('claims', '')[:80]!r}{_C_RST}", flush=True)

        print(f"{_C_SEND}[EMBED    →] model={self.llm.embedding_model}  text={len(summary)} chars{_C_RST}", flush=True)
        embedding, embed_usage = await self.llm.embed(summary)
        self.state.add_usage("embedding", embed_usage)
        print(f"{_C_RECV}[EMBED    ←] dims={len(embedding)}  tokens={embed_usage['total_tokens']}{_C_RST}", flush=True)

        store_full = self.cfg["memory"].get("store_full_message", True)
        full_text = f"{combined_content}{image_context}{doc_context}" if store_full else ""
        self.memory.add(
            text=summary,
            embedding=embedding,
            metadata={
                "source_message_id": group_ids[0],
                "group_message_ids": ",".join(group_ids),
                "author_id": str(primary.author.id),
                "author_name": primary.author.display_name,
                "author_discord_name": primary.author.name,
                "channel": primary.channel.name,
                "channel_id": str(primary.channel.id),
                "timestamp": primary.created_at.isoformat(),
                "year": int(ingest_year),
                "entry_type": "message_chain",
                "has_image": bool(all_images),
                "full_message_text": full_text,
                "original_excerpt": combined_content[:500],
                "entities": tags.get("entities", ""),
                "action_type": tags.get("action_type", ""),
                "claims": tags.get("claims", ""),
            },
        )
        if doc_parts:
            doc_chunks = await self._store_document_chunks(
                doc_parts=doc_parts,
                source_message_id=group_ids[0],
                group_message_ids=",".join(group_ids),
                author_id=str(primary.author.id),
                author_name=primary.author.display_name,
                author_discord_name=primary.author.name,
                channel_name=primary.channel.name,
                channel_id=str(primary.channel.id),
                timestamp=primary.created_at.isoformat(),
                year=int(ingest_year),
            )
            if doc_chunks:
                print(f"{_C_OK}[DOC STORE✓] {doc_chunks} document chunks saved{_C_RST}", flush=True)
        self.state.bump("messages_archived")
        self.flog.log_archived(
            group_ids[0], primary.author.display_name,
            primary.channel.name, len(combined_content),
            int(ingest_year),
        )
        print(f"{_C_OK}[STORE    ✓] chain memory saved ({len(messages)} msgs)  total={self.memory.count()}{_C_RST}", flush=True)
        print(f"{_C_HDR}{'─'*70}{_C_RST}", flush=True)
        return True

    async def _gather_action_chain(self, target_msg: discord.Message) -> list[discord.Message]:
        """Return the full run of consecutive same-author messages around target_msg.

        Walks backward and forward through channel history, stopping as soon as
        another author posts or the gap exceeds chain_delay_seconds.
        """
        from datetime import timedelta
        threshold = timedelta(seconds=self._chain_delay)

        try:
            before_msgs = [
                m async for m in target_msg.channel.history(
                    limit=10, before=target_msg, oldest_first=False
                )
            ]
            after_msgs = [
                m async for m in target_msg.channel.history(
                    limit=10, after=target_msg, oldest_first=True
                )
            ]
        except (discord.Forbidden, discord.HTTPException):
            return [target_msg]

        chain: list[discord.Message] = [target_msg]

        # Expand backward (before_msgs[0] is the message immediately before target)
        for m in before_msgs:
            if m.author.id != target_msg.author.id:
                break
            if chain[0].created_at - m.created_at > threshold:
                break
            chain.insert(0, m)

        # Expand forward
        for m in after_msgs:
            if m.author.id != target_msg.author.id:
                break
            if m.created_at - chain[-1].created_at > threshold:
                break
            chain.append(m)

        return chain

    # ------------------------------------------------------------------
    # Send helper
    # ------------------------------------------------------------------
    async def _send_chunked_reply(self, message: discord.Message, text: str):
        for i in range(0, len(text), 1900):
            chunk = text[i:i + 1900]
            if i == 0:
                await message.reply(chunk)
            else:
                await message.channel.send(chunk)

    # ------------------------------------------------------------------
    # Public commands (small)
    # ------------------------------------------------------------------
    async def _cmd_optout(self, message: discord.Message):
        added = self.optouts.add(str(message.author.id), message.author.display_name)
        if added:
            await message.reply(
                "✅ Got it — your future messages won't be archived.\n"
                f"Use `{self.prefix} optin` to reverse this. To remove memories already archived, "
                f"ask an admin to run `{self.prefix} void` on you."
            )
        else:
            await message.reply("You're already opted out.")

    async def _cmd_optin(self, message: discord.Message):
        removed = self.optouts.remove(str(message.author.id))
        if removed:
            await message.reply("✅ Welcome back — your messages will be archived again.")
        else:
            await message.reply("You weren't opted out.")

    async def _cmd_whoami(self, message: discord.Message):
        member = message.author
        if not isinstance(member, discord.Member):
            await message.reply("Couldn't determine your member status in this guild.")
            return
        roles = [r.name for r in member.roles if r.name != "@everyone"]
        role_set_lower = {r.lower() for r in roles}
        admin_match = role_set_lower & self._admin_roles
        gm_match = role_set_lower & self._gm_roles
        perms = message.channel.permissions_for(member)
        manage_guild = perms.manage_guild

        is_admin = self._is_admin_user(message)
        is_gm = self._is_gm_user(member)

        body = (
            f"```\n"
            f"You are: {member.display_name} ({member.id})\n"
            f"Your roles: {', '.join(roles) if roles else '(none)'}\n"
            f"\n"
            f"admin_roles in config: {sorted(self._admin_roles) or '(none)'}\n"
            f"gm_roles in config:    {sorted(self._gm_roles) or '(none)'}\n"
            f"\n"
            f"Manage Server permission: {'YES' if manage_guild else 'no'}\n"
            f"Matched admin role(s):    {sorted(admin_match) or '(none)'}\n"
            f"Matched GM role(s):       {sorted(gm_match) or '(none)'}\n"
            f"\n"
            f"Bot will let you:\n"
            f"  Run admin commands:     {'YES' if is_admin else 'no'}\n"
            f"  Use !DNC chatuc:        {'YES' if is_admin else 'no'}\n"
            f"  Revise GM rulings:      {'YES' if is_gm else 'no'}\n"
            f"```"
        )
        await message.reply(body)

    # ------------------------------------------------------------------
    # Espionage system
    # ------------------------------------------------------------------

    async def _check_espionage(self, message: discord.Message) -> None:
        """Roll for interception by each spy watching the poster's TAG."""
        tag = parse_tag(message.author.display_name)
        if not tag:
            return
        spies = self.espionage.get_spies_for_tag(tag)
        if not spies:
            return

        base = float(self._espionage_cfg.get("base_chance", 30))
        modifiers: dict = self._espionage_cfg.get("base_modifiers") or {}
        tag_mod = float(modifiers.get(tag, 0))
        cspy_mult = float(self._espionage_cfg.get("counterspy_multiplier", 0.25))
        is_cspy = self.espionage.is_counterspy(str(message.author.id))

        effective = (base + tag_mod) * (cspy_mult if is_cspy else 1.0)
        effective = max(0.0, min(100.0, effective))

        content = (message.content or "").strip()
        attachments = message.attachments

        for spy_uid in spies:
            if random.uniform(0, 100) > effective:
                continue
            try:
                spy_member = message.guild.get_member(int(spy_uid))
                if spy_member is None:
                    continue
                dm = await spy_member.create_dm()
                body_lines = [
                    f"🕵️ **Intelligence Intercept** — Year {self.state.current_year}",
                    f"Your operatives have intercepted a communication from **{tag}**"
                    f" ({message.author.display_name}) in **#{message.channel.name}**:",
                    "",
                ]
                if content:
                    body_lines.append(f">>> {content}")
                if attachments:
                    urls = "  ".join(a.url for a in attachments)
                    body_lines.append(
                        f"\n⚠️ This message included {len(attachments)} attachment(s): {urls}"
                    )
                await dm.send("\n".join(body_lines))
            except (discord.Forbidden, discord.HTTPException):
                pass

    async def _cmd_spy(self, message: discord.Message, arg: str) -> None:
        if not self._is_espionage_command_channel(message.channel):
            await message.reply(
                f"⚠️ Intelligence commands must be used in the designated channel."
            )
            return

        tag = arg.strip().upper()
        if not tag:
            await message.reply(
                f"Usage: `{self.prefix} spy <TAG>` — begin monitoring a nation's secret communications."
            )
            return

        target = self._find_member_by_tag(message.guild, tag)
        if target is None:
            await message.reply(f"⚠️ No player found with tag `{tag}`.")
            return
        if target.id == message.author.id:
            await message.reply("⚠️ You cannot run intelligence operations against yourself.")
            return

        user_id = str(message.author.id)
        entry = self.espionage.get_entry(user_id)
        max_slots = 1 if entry["counterspy"] else 3

        result = self.espionage.add_target(user_id, tag, max_slots)
        if result == "duplicate":
            await message.reply(
                f"Your intelligence services are already monitoring **{tag}**."
            )
        elif result == "full":
            limit_str = "1 (counterspy active)" if entry["counterspy"] else "3"
            await message.reply(
                f"⚠️ Intelligence capacity reached ({limit_str} targets). "
                f"Use `{self.prefix} unspy <TAG>` to stand down an operation first."
            )
        else:
            current = self.espionage.list_targets(user_id)
            await message.reply(
                f"🕵️ Intelligence operations against **{tag}** are now active. "
                f"({len(current)}/{max_slots} slots used)"
            )

    async def _cmd_counterspy(self, message: discord.Message) -> None:
        if not self._is_espionage_command_channel(message.channel):
            await message.reply(
                f"⚠️ Intelligence commands must be used in the designated channel."
            )
            return

        user_id = str(message.author.id)
        entry = self.espionage.get_entry(user_id)

        if entry["counterspy"]:
            # Toggle off
            self.espionage.set_counterspy(user_id, False)
            await message.reply(
                "🔓 Counterintelligence protocols **deactivated**. "
                "Your communications are no longer encrypted. "
                "You may now run operations against up to **3** targets."
            )
        else:
            # Toggle on — trim spy slots to 1
            dropped = self.espionage.trim_targets(user_id, 1)
            self.espionage.set_counterspy(user_id, True)
            lines = [
                "🔐 Counterintelligence protocols **activated**. "
                "Your communications are now heavily encrypted. "
                "You are limited to **1** intelligence target."
            ]
            if dropped:
                lines.append(
                    f"⚠️ Operations against **{', '.join(dropped)}** have been stood down "
                    f"to free resources for counterintelligence."
                )
            await message.reply("\n".join(lines))

    async def _cmd_unspy(self, message: discord.Message, arg: str) -> None:
        if not self._is_espionage_command_channel(message.channel):
            await message.reply(
                f"⚠️ Intelligence commands must be used in the designated channel."
            )
            return

        tag = arg.strip().upper()
        if not tag:
            await message.reply(
                f"Usage: `{self.prefix} unspy <TAG>` — stand down operations against a target."
            )
            return

        user_id = str(message.author.id)
        if self.espionage.remove_target(user_id, tag):
            await message.reply(
                f"🔍 Intelligence operations against **{tag}** have been stood down."
            )
        else:
            await message.reply(
                f"⚠️ You are not currently running operations against `{tag}`."
            )

    async def _cmd_spylist(self, message: discord.Message) -> None:
        if not self._is_espionage_command_channel(message.channel):
            await message.reply(
                f"⚠️ Intelligence commands must be used in the designated channel."
            )
            return

        user_id = str(message.author.id)
        entry = self.espionage.get_entry(user_id)
        targets = entry["targets"]
        cspy = entry["counterspy"]
        max_slots = 1 if cspy else 3

        if not targets and not cspy:
            await message.reply("You have no active intelligence operations.")
            return

        lines = [f"**Intelligence Operations** ({len(targets)}/{max_slots} targets active)"]
        for t in targets:
            lines.append(f"• **{t}** — monitoring")
        if cspy:
            lines.append("\n🔐 Counterintelligence: **ACTIVE** (communications encrypted)")
        else:
            lines.append("\n🔓 Counterintelligence: inactive")

        # Base interception chance as a rough guide
        base = float(self._espionage_cfg.get("base_chance", 30))
        cspy_mult = float(self._espionage_cfg.get("counterspy_multiplier", 0.25))
        own_tag = parse_tag(message.author.display_name) or "?"
        if cspy:
            lines.append(
                f"📊 Others spying on **{own_tag}**: ~{base * cspy_mult:.0f}% intercept chance (reduced)"
            )
        else:
            lines.append(
                f"📊 Others spying on **{own_tag}**: ~{base:.0f}% base intercept chance"
            )

        await message.reply("\n".join(lines))

    # ------------------------------------------------------------------
    # Admin: ingest backfill
    # ------------------------------------------------------------------
    async def _cmd_ingest(self, message: discord.Message, arg: str):
        target_channel = message.channel
        rest_arg = arg.strip()

        # Strip optional trailing "year <N>" override before any other parsing
        year_override = None
        year_match = re.search(r'\byear\s+(\d+)$', rest_arg, re.IGNORECASE)
        if year_match:
            year_override = int(year_match.group(1))
            rest_arg = rest_arg[:year_match.start()].strip()

        # 1. Check for single message link or ID
        link_match = MSG_LINK_RE.match(rest_arg)
        id_match = ID_RE.match(rest_arg)

        if link_match or id_match:
            await self._ingest_single_target(message, rest_arg, link_match, id_match, year_override=year_override)
            return

        # 2. Parse the channel tag if present
        ch_match = re.match(r"<#(\d+)>\s+(.+)", rest_arg)
        if ch_match:
            ch = message.guild.get_channel(int(ch_match.group(1)))
            if ch is None:
                await message.reply("⚠️ Couldn't find that channel.")
                return
            target_channel = ch
            rest_arg = ch_match.group(2).strip()

        # 2. Figure out if we are using a count (N) or a date range
        parts = rest_arg.split()
        limit = None
        after_dt = None
        before_dt = None

        if len(parts) == 1 and parts[0].isdigit():
            # Legacy behavior: just a number
            limit = int(parts[0])
            if limit < 1 or limit > 5000:
                await message.reply("Pick a number between 1 and 5000.")
                return
        elif len(parts) >= 1:
            # Date range behavior
            after_dt = self._parse_date(parts[0])
            if not after_dt:
                await message.reply(
                    f"Usage: `{self.prefix} ingest [#channel] <N>` OR \n"
                    f"`{self.prefix} ingest [#channel] <start-date> [end-date]` (Format: MM-DD-YYYY)"
                )
                return
            
            # If an end date was provided, parse it. Otherwise, default to now.
            if len(parts) >= 2:
                before_dt = self._parse_date(parts[1])
                if not before_dt:
                    await message.reply("⚠️ Could not parse the end date.")
                    return

        if not self._is_scanned(target_channel):
            await message.reply(f"⚠️ #{target_channel.name} is in the ignored or non-whitelisted list.")
            return

        # 3. Setup the status message
        year_suffix = f" — stamping as year {year_override}" if year_override is not None else ""
        if limit:
            status = await message.reply(f"📜 Backfilling up to {limit} messages from #{target_channel.name}{year_suffix}…")
        else:
            end_str = (before_dt or datetime.now(timezone.utc)).strftime('%Y-%m-%d')
            status = await message.reply(f"📜 Backfilling #{target_channel.name} from {after_dt.strftime('%Y-%m-%d')} to {end_str}{year_suffix}…")

        scanned = skipped_short = skipped_dupe = skipped_bot = 0
        skipped_optout = archived = non_lore = errors = 0

        try:
            # 4. Fetch the history using the dynamic parameters
            fetch_limit = limit if limit else 5000 # hard cap for safety if using dates
            
            history_iter = target_channel.history(
                limit=fetch_limit,
                after=after_dt,
                before=before_dt,
                oldest_first=bool(after_dt) # Read forward in time if using a date range
            )
            
            history = [m async for m in history_iter]
            
            # If we used a limit (no dates), we grabbed the newest N messages backwards. 
            # Reverse them to process chronologically.
            if not after_dt:
                history.reverse()

            for msg in history:
                scanned += 1
                if msg.author.bot:
                    skipped_bot += 1; continue
                if len((msg.content or "").strip()) < self._get_min_length(msg):
                    skipped_short += 1; continue
                if self.optouts.is_opted_out(str(msg.author.id)):
                    skipped_optout += 1; continue
                if self.memory.has_source_message(str(msg.id)):
                    skipped_dupe += 1; continue
                try:
                    if await self._ingest_message(msg, year_override=year_override): archived += 1
                    else: non_lore += 1
                except Exception:
                    log.exception("Backfill ingest failed for %s", msg.id)
                    errors += 1
                if archived and archived % 10 == 0:
                    try:
                        await status.edit(content=(
                            f"📜 Backfilling #{target_channel.name}…\n"
                            f"Scanned {scanned}/{len(history)} • archived **{archived}** • "
                            f"non-lore {non_lore} • short {skipped_short} • "
                            f"dupe {skipped_dupe} • optout {skipped_optout}"
                        ))
                    except discord.HTTPException:
                        pass
                await asyncio.sleep(0.05)
        except discord.Forbidden:
            await status.edit(content="⚠️ No permission to read that channel's history.")
            return
        except Exception as e:
            log.exception("Backfill failed")
            await status.edit(content=f"⚠️ Backfill failed: `{e}`")
            return

        await status.edit(content=(
            f"✅ Backfill complete for #{target_channel.name}\n```\n"
            f"scanned         {scanned}\n"
            f"archived        {archived}\n"
            f"non-lore        {non_lore}\n"
            f"too short       {skipped_short}\n"
            f"already in db   {skipped_dupe}\n"
            f"opted out       {skipped_optout}\n"
            f"bot messages    {skipped_bot}\n"
            f"errors          {errors}\n"
            f"total memories  {self.memory.count()}\n```"
        ))

    async def _ingest_single_target(self, message: discord.Message, rest_arg: str,
                                   link_match, id_match, year_override: int | None = None):
        """Helper to fetch and ingest a single message by link or ID."""
        try:
            if link_match:
                # It's a Discord message link
                channel_id, msg_id = link_match.group(2), link_match.group(3)
                fetch_channel = message.guild.get_channel(int(channel_id))
            else:
                # It's a raw ID (assume current channel)
                msg_id = rest_arg
                fetch_channel = message.channel

            if not fetch_channel:
                await message.reply("⚠️ Couldn't find the channel for that message.")
                return

            target_msg = await fetch_channel.fetch_message(int(msg_id))
            status = await message.reply("📜 Ingesting specific message...")

            # Check for duplicates, opt-outs, and length
            if self.memory.has_source_message(str(target_msg.id)):
                await status.edit(content="⚠️ That message is already in the archives.")
                return
            if self.optouts.is_opted_out(str(target_msg.author.id)):
                await status.edit(content="⚠️ That user is opted out.")
                return
            if len((target_msg.content or "").strip()) < self._get_min_length(target_msg):
                await status.edit(content="⚠️ Message is too short to ingest.")
                return

            # Attempt ingestion
            if await self._ingest_message(target_msg, year_override=year_override):
                year_note = f" (year {year_override})" if year_override is not None else ""
                await status.edit(content=f"✅ Successfully archived message `{msg_id}`{year_note}!")
            else:
                await status.edit(content=f"⚠️ Message scanned but filtered (e.g., flagged as NO_LORE).")

        except discord.NotFound:
            await message.reply("⚠️ Couldn't find that message. Make sure the ID/link is correct.")
        except discord.Forbidden:
            await message.reply("⚠️ I don't have permission to read that channel.")
        except Exception as e:
            log.exception("Single message ingest failed")
            await message.reply(f"⚠️ Failed: `{e}`")

    # ------------------------------------------------------------------
    # Admin: void / unvoid
    # ------------------------------------------------------------------
    async def _cmd_void(self, message: discord.Message, arg: str):
        arg = arg.strip()
        if not arg:
            if message.reference and message.reference.message_id:
                await self._void_message(message, str(message.reference.message_id))
                return
            await message.reply(
                f"Usage: `{self.prefix} void <message-link>` "
                f"or `{self.prefix} void <@user>` "
                f"or `{self.prefix} void msg <id>` / `{self.prefix} void user <id>`\n"
                f"You can also reply to a message and run `{self.prefix} void` with no arguments."
            )
            return

        explicit = arg.split(maxsplit=1)
        if len(explicit) == 2 and explicit[0].lower() in ("msg", "user"):
            kind = "message" if explicit[0].lower() == "msg" else "user"
            target_id = explicit[1].strip()
            if not ID_RE.match(target_id):
                await message.reply("⚠️ That doesn't look like a valid Discord ID.")
                return
            if kind == "message":
                await self._void_message(message, target_id)
            else:
                await self._void_user(message, target_id)
            return

        kind, info = parse_void_target(arg)
        if kind == "message_link":
            await self._void_message(message, info["message_id"])
        elif kind == "mention":
            await self._void_user(message, info["user_id"])
        elif kind == "snowflake":
            try:
                user = await self.fetch_user(int(info["id"]))
                await self._void_user(message, info["id"], known_user=user)
                return
            except discord.NotFound:
                pass
            except Exception:
                pass
            await self._void_message(message, info["id"])
        else:
            await message.reply("⚠️ Couldn't parse that. Paste a message link, mention a user, or use `void msg <id>` / `void user <id>`.")

    async def _void_message(self, invocation: discord.Message, message_id: str):
        existing = self.memory.get_by_source_message(message_id)
        if not existing:
            await invocation.reply(f"⚠️ No memory found for message ID `{message_id}`.")
            return
        author_name = existing["metadata"].get("author_name", "?")
        group_id = self.memory.void_by_source_message(
            message_id, voided_by=str(invocation.author)
        )
        if not group_id:
            await invocation.reply(f"⚠️ Void failed for `{message_id}`.")
            return
        self.state.bump("voids_executed")
        self.flog.log_void_msg(message_id, author_name, str(invocation.author), group_id)
        await invocation.reply(
            f"🗑️ Voided memory from message `{message_id}` (author: {author_name}). "
            f"Recoverable for {self.cfg['memory']['void_retention_days']} days."
        )

    async def _void_user(self, invocation, user_id: str, known_user=None):
        memories = self.memory.get_by_author(user_id)
        if not memories:
            await invocation.reply(f"⚠️ No memories found for user ID `{user_id}`.")
            return
        user_name = (
            known_user.display_name if known_user
            else memories[0]["metadata"].get("author_name", user_id)
        )
        confirm = await invocation.reply(
            f"⚠️ This will void **{len(memories)}** memories authored by "
            f"**{user_name}** (`{user_id}`). React ✅ within 30 seconds to confirm."
        )
        try:
            await confirm.add_reaction("✅")
        except discord.HTTPException:
            pass

        def check(reaction, user):
            return (
                reaction.message.id == confirm.id
                and user.id == invocation.author.id
                and str(reaction.emoji) == "✅"
            )
        try:
            await self.wait_for("reaction_add", timeout=30.0, check=check)
        except asyncio.TimeoutError:
            await confirm.edit(content="⏱️ Void cancelled — no confirmation received.")
            return

        group_id, count = self.memory.void_by_author(user_id, voided_by=str(invocation.author))
        if not group_id:
            await confirm.edit(content="⚠️ Void failed.")
            return
        self.state.bump("voids_executed")
        self.flog.log_void_user(user_id, user_name, count, str(invocation.author), group_id)
        await confirm.edit(content=(
            f"🗑️ Voided {count} memories by **{user_name}** (`{user_id}`). "
            f"Recoverable for {self.cfg['memory']['void_retention_days']} days."
        ))

    async def _cmd_unvoid(self, message: discord.Message, arg: str):
        arg = arg.strip()
        if not arg:
            await message.reply(f"Usage: `{self.prefix} unvoid <message-id>` or `{self.prefix} unvoid <user-id>`")
            return

        m = MENTION_RE.match(arg)
        if m:
            target_id = m.group(1)
        elif ID_RE.match(arg):
            target_id = arg
        else:
            link = MSG_LINK_RE.match(arg)
            if link:
                target_id = link.group(3)
            else:
                await message.reply("⚠️ Invalid ID or message link.")
                return

        msg_tombs = self.memory.find_voided_by_source_message(target_id)
        user_tombs = self.memory.find_voided_by_author(target_id)
        tombs = msg_tombs or user_tombs
        if not tombs:
            await message.reply(
                f"⚠️ No voided memories found for `{target_id}`. "
                f"Past the {self.cfg['memory']['void_retention_days']}-day window, or wrong ID."
            )
            return

        groups = {t.get("void_group_id") for t in tombs if t.get("void_group_id")}
        total = 0
        for g in groups:
            total += self.memory.unvoid_group(g)
        kind = "message" if msg_tombs else "user"
        self.state.bump("unvoids_executed")
        self.flog.log_unvoid(f"{kind}={target_id}", total, str(message.author),
                             ",".join(filter(None, groups)))
        await message.reply(f"♻️ Restored {total} voided memories ({kind} `{target_id}`).")

    # ------------------------------------------------------------------
    # Admin: year / purge / export
    # ------------------------------------------------------------------
    async def _cmd_yearset(self, message: discord.Message, arg: str):
        try:
            year = int(arg.strip())
        except ValueError:
            await message.reply(f"Usage: `{self.prefix} yearset <year>`")
            return
        old = self.state.current_year
        self.state.set_year(year, mark_rollover=True)
        await message.reply(f"📅 Year set to **{year}** (was {old}). No announcement.")

    async def _cmd_yearroll(self, message: discord.Message):
        new_year = self.state.current_year + 1
        self.state.set_year(new_year, mark_rollover=True)
        await self._announce_year_roll(new_year)
        await message.reply(f"📅 Year rolled to **{new_year}** and announcement posted.")

    async def _cmd_purge(self, message: discord.Message, arg: str):
        parts = arg.strip().split(maxsplit=1)
        if len(parts) != 2 or parts[0].lower() != "year":
            await message.reply(f"Usage: `{self.prefix} purge year <year>`")
            return
        try:
            year = int(parts[1])
        except ValueError:
            await message.reply("⚠️ Invalid year.")
            return

        all_memories = self.memory.get_all()
        targets = [m for m in all_memories if m["metadata"].get("year") == year]
        if not targets:
            await message.reply(f"⚠️ No memories found for year **{year}**.")
            return

        confirm = await message.reply(
            f"⚠️ This will VOID (recoverable for "
            f"{self.cfg['memory']['void_retention_days']} days) "
            f"**{len(targets)}** memories from year **{year}**. "
            f"React ✅ within 30 seconds to confirm."
        )
        try:
            await confirm.add_reaction("✅")
        except discord.HTTPException:
            pass

        def check(reaction, user):
            return (
                reaction.message.id == confirm.id
                and user.id == message.author.id
                and str(reaction.emoji) == "✅"
            )
        try:
            await self.wait_for("reaction_add", timeout=30.0, check=check)
        except asyncio.TimeoutError:
            await confirm.edit(content="⏱️ Purge cancelled.")
            return

        ids = [m["id"] for m in targets]
        count = self.memory.void(
            memory_ids=ids,
            void_reason=f"purge year={year}",
            voided_by=str(message.author),
        )
        self.state.bump("voids_executed")
        await confirm.edit(content=f"🗑️ Voided {count} memories from year {year} (recoverable).")

    async def _cmd_export(self, message: discord.Message):
        os.makedirs("exports", exist_ok=True)
        ts = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
        json_path = os.path.join("exports", f"memories-{ts}.json")
        md_path = os.path.join("exports", f"memories-{ts}.md")

        all_mem = self.memory.get_all()
        all_mem.sort(key=lambda m: (
            int(m["metadata"].get("year", 0) or 0),
            m["metadata"].get("timestamp", "") or "",
        ))

        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(all_mem, f, indent=2, ensure_ascii=False)

        with open(md_path, "w", encoding="utf-8") as f:
            f.write(f"# DNC Memory Archive Export — {ts}\n\n")
            f.write(f"Total memories: {len(all_mem)}\n\n---\n\n")
            for m in all_mem:
                meta = m["metadata"]
                f.write(f"## Year {meta.get('year', '?')} "
                        f"[{meta.get('entry_type', 'message')}]\n")
                f.write(f"- **Author:** {meta.get('author_name', '?')} (`{meta.get('author_id','?')}`)\n")
                f.write(f"- **Channel:** #{meta.get('channel', '?')}\n")
                f.write(f"- **Posted:** {meta.get('timestamp', '?')}\n")
                f.write(f"- **Source message:** `{meta.get('source_message_id', '?')}`\n\n")
                f.write(f"**Summary:**\n\n{m['text']}\n\n")
                full = meta.get("full_message_text") or meta.get("original_excerpt", "")
                if full:
                    f.write(f"<details><summary>Full text</summary>\n\n```\n{full}\n```\n</details>\n\n")
                f.write("---\n\n")

        await message.reply(
            f"✅ Exported {len(all_mem)} memories to:\n`{json_path}`\n`{md_path}`"
        )

    # ------------------------------------------------------------------
    # Admin: stats
    # ------------------------------------------------------------------
    async def _cmd_stats(self, message: discord.Message, arg: str):
        if arg.strip().lower() == "reset":
            self.state.reset_stats()
            await message.reply("✅ Stats counters reset.")
            return

        s = self.state.stats_snapshot()
        started = self.state.started_at or "?"
        uptime_str = "?"
        try:
            started_dt = datetime.fromisoformat(started)
            delta = datetime.now(timezone.utc) - started_dt
            hours = delta.total_seconds() / 3600
            uptime_str = f"{hours:.1f}h ({delta.days}d)"
        except Exception:
            pass

        all_tokens = (
            s.get("archivist_prompt_tokens", 0) + s.get("archivist_completion_tokens", 0)
            + s.get("query_prompt_tokens", 0) + s.get("query_completion_tokens", 0)
            + s.get("chat_prompt_tokens", 0) + s.get("chat_completion_tokens", 0)
            + s.get("chatuc_prompt_tokens", 0) + s.get("chatuc_completion_tokens", 0)
            + s.get("gm_prompt_tokens", 0) + s.get("gm_completion_tokens", 0)
            + s.get("war_prompt_tokens", 0) + s.get("war_completion_tokens", 0)
            + s.get("embedding_tokens", 0)
        )

        body = (
            f"```\n"
            f"=== DNC stats ===\n"
            f"current year       {self.state.current_year}\n"
            f"total memories     {self.memory.count()}\n"
            f"opt-outs           {self.optouts.count()}\n"
            f"bot started        {started}\n"
            f"uptime             {uptime_str}\n"
            f"counters since     {s.get('stats_reset_at') or 'first start'}\n"
            f"\n"
            f"--- messages ---\n"
            f"seen                  {s.get('messages_seen', 0):>8}\n"
            f"filtered (local)      {s.get('messages_filtered_local', 0):>8}\n"
            f"sent to archivist     {s.get('messages_sent_to_archivist', 0):>8}\n"
            f"archived              {s.get('messages_archived', 0):>8}\n"
            f"filtered (NO_LORE)    {s.get('messages_filtered_non_lore', 0):>8}\n"
            f"\n"
            f"--- commands ---\n"
            f"queries answered      {s.get('queries_answered', 0):>8}\n"
            f"chats answered        {s.get('chats_answered', 0):>8}\n"
            f"chatuc answered       {s.get('chatuc_answered', 0):>8}\n"
            f"GM rulings made       {s.get('gm_rulings_made', 0):>8}\n"
            f"GM rulings revised    {s.get('gm_rulings_revised', 0):>8}\n"
            f"wars started          {s.get('wars_started', 0):>8}\n"
            f"war moves adjudicated {s.get('war_moves_adjudicated', 0):>8}\n"
            f"war AI moves          {s.get('war_npc_moves', 0):>8}\n"
            f"wars ended            {s.get('wars_ended', 0):>8}\n"
            f"voids                 {s.get('voids_executed', 0):>8}\n"
            f"unvoids               {s.get('unvoids_executed', 0):>8}\n"
            f"\n"
            f"--- tokens (prompt / completion) ---\n"
            f"archivist  {s.get('archivist_prompt_tokens', 0):>10} / {s.get('archivist_completion_tokens', 0):<10}\n"
            f"queries    {s.get('query_prompt_tokens', 0):>10} / {s.get('query_completion_tokens', 0):<10}\n"
            f"chat       {s.get('chat_prompt_tokens', 0):>10} / {s.get('chat_completion_tokens', 0):<10}\n"
            f"chatuc     {s.get('chatuc_prompt_tokens', 0):>10} / {s.get('chatuc_completion_tokens', 0):<10}\n"
            f"gm         {s.get('gm_prompt_tokens', 0):>10} / {s.get('gm_completion_tokens', 0):<10}\n"
            f"war        {s.get('war_prompt_tokens', 0):>10} / {s.get('war_completion_tokens', 0):<10}\n"
            f"embeddings {s.get('embedding_tokens', 0):>10} (total)\n"
            f"GRAND TOTAL TOKENS    {all_tokens:>8}\n"
            f"```"
        )
        await message.reply(body)

    # ------------------------------------------------------------------
    # ------------------------------------------------------------------
    # Admin: chat blacklist
    # ------------------------------------------------------------------
    async def _cmd_chatban(self, message: discord.Message, arg: str):
        arg = arg.strip()
        if not arg:
            await message.reply(f"Usage: `{self.prefix} chatban <@user|ID>`")
            return
        m = MENTION_RE.match(arg)
        user_id = m.group(1) if m else arg
        if not ID_RE.match(user_id):
            await message.reply("⚠️ Couldn't parse a user ID from that argument.")
            return
        display_name = user_id
        try:
            user = await self.fetch_user(int(user_id))
            display_name = user.display_name
        except Exception:
            pass
        added = self.chat_blacklist.add(user_id, display_name)
        if added:
            await message.reply(f"🚫 {display_name} (`{user_id}`) added to the chat blacklist.")
        else:
            await message.reply(f"ℹ️ {display_name} (`{user_id}`) is already on the chat blacklist.")

    async def _cmd_chatunban(self, message: discord.Message, arg: str):
        arg = arg.strip()
        if not arg:
            await message.reply(f"Usage: `{self.prefix} chatunban <@user|ID>`")
            return
        m = MENTION_RE.match(arg)
        user_id = m.group(1) if m else arg
        if not ID_RE.match(user_id):
            await message.reply("⚠️ Couldn't parse a user ID from that argument.")
            return
        removed = self.chat_blacklist.remove(user_id)
        if removed:
            await message.reply(f"✅ `{user_id}` removed from the chat blacklist.")
        else:
            await message.reply(f"ℹ️ `{user_id}` was not on the chat blacklist.")

    # Admin: channels diagnostic
    # ------------------------------------------------------------------
    async def _cmd_channels(self, message: discord.Message):
        guild = message.guild
        actual_by_norm = {
            normalize_channel_name(ch.name): ch for ch in guild.text_channels
        }

        def render_section(title: str, configured: list) -> list:
            lines = [f"=== {title} ==="]
            if not configured:
                lines.append("  (none configured)")
                return lines
            for entry in configured:
                norm = normalize_channel_name(entry)
                actual = actual_by_norm.get(norm)
                if actual:
                    lines.append(f"  ✓ {entry!r:<40s} → matches #{actual.name}")
                else:
                    extra = ""
                    if any(c.isspace() for c in entry):
                        extra = "  [contains whitespace]"
                    lines.append(f"  ✗ {entry!r:<40s} → no match{extra}")
            return lines

        lines = []
        lines += render_section("scan_channels (whitelist; empty=all non-ignored)", self._scan_raw)
        lines.append("")
        lines += render_section("ignored_channels", self._ignored_raw)
        lines.append("")
        lines += render_section("admin_channels", self._admin_raw)
        lines.append("")
        lines.append(f"=== Guild has {len(guild.text_channels)} text channels ===")

        body = "\n".join(lines)
        chunks = []
        cur = ""
        for line in body.split("\n"):
            if len(cur) + len(line) + 1 > 1900:
                chunks.append(cur); cur = line
            else:
                cur = cur + "\n" + line if cur else line
        if cur:
            chunks.append(cur)
        first = True
        for c in chunks:
            wrapped = f"```\n{c}\n```"
            if first:
                await message.reply(wrapped); first = False
            else:
                await message.channel.send(wrapped)

    async def _cmd_reloadprompts(self, message: discord.Message):
        results = self.prompts.reload()
        lines = ["**Prompt reload results:**"]
        for name, status in results.items():
            icon = "✓" if status == "ok" else "✗"
            lines.append(f"`{icon} {name}` — {status}")
        await message.reply("\n".join(lines))

    # ------------------------------------------------------------------
    def _get_map_channel(self) -> Optional[discord.TextChannel]:
        name = normalize_channel_name(
            self.cfg.get("spatial_mapping", {}).get("output_channel", "") or ""
        )
        if not name:
            return None
        for guild in self.guilds:
            for ch in guild.text_channels:
                if normalize_channel_name(ch.name) == name:
                    return ch
        return None

    def _province_render_args(self) -> dict:
        """Build keyword args for map_renderer calls that include province ownership."""
        return dict(province_ownership=self.province_store.get_ownership())

    async def _prewarm_map_cache(self) -> None:
        """Render the world map once at startup so the first !DNC map is hot.

        Runs as a background task off on_ready. If the disk-persisted entry's
        state-key still matches the current state, the cache reuses it
        without invoking the renderer.
        """
        try:
            for guild in self.guilds:
                self.map_store.rebuild(map_sched.scan_members(guild.members))
            _png, key, hit = await self._render_world_via_cache()
            log.info("Map cache prewarm: %s key=%s", "hit" if hit else "miss", key[:8])
        except Exception:
            log.exception("Map cache prewarm failed")

    async def _render_world_via_cache(self) -> tuple[bytes, str, bool]:
        """Render the world map through the cache. Returns (png, key, was_hit)."""
        occupation = self.map_store.get_occupation()
        ownership = self.province_store.get_ownership()

        async def _render() -> bytes:
            return await map_renderer.render_world_map(
                occupation, province_ownership=ownership, year=self.state.current_year,
            )

        png, key, hit = await self.map_cache.get_or_render(
            occupation, ownership, self.state.current_year, _render,
        )
        log.info("Map cache: %s key=%s", "hit" if hit else "miss", key[:8])
        return png, key, hit

    async def _render_and_post_map(
        self, channel: Optional[discord.TextChannel] = None
    ) -> Optional[discord.Message]:
        target = channel or self._get_map_channel()
        if target is None:
            log.warning("spatial_mapping: output_channel not found, skipping render")
            return None
        try:
            img_bytes, _key, _hit = await self._render_world_via_cache()
        except Exception as e:
            log.error("Map render failed: %s", e, exc_info=True)
            return None
        try:
            return await target.send(
                file=discord.File(fp=io.BytesIO(img_bytes), filename="dnc_map.png")
            )
        except discord.Forbidden:
            log.warning("Map: no permission to post in #%s", target.name)
            return None

    async def _cmd_map(self, message: discord.Message) -> None:
        # Defensive resync — a missed on_member_update would otherwise leave
        # stale occupation, and the cache key follows whatever state we
        # actually hold. Rebuild is idempotent when nothing changed.
        for guild in self.guilds:
            self.map_store.rebuild(map_sched.scan_members(guild.members))
        async with message.channel.typing():
            try:
                img_bytes, _key, _hit = await self._render_world_via_cache()
            except Exception as e:
                await message.reply(f"Map render failed: {e}")
                return
        await message.reply(
            file=discord.File(fp=io.BytesIO(img_bytes), filename="dnc_map.png")
        )

    async def _cmd_mapzoom(self, message: discord.Message, tag: str) -> None:
        tag = tag.strip().upper()
        if not re.match(r"^[A-Z]{2,4}$", tag):
            await message.reply(
                "Please provide a valid ISO country tag (2–4 uppercase letters), "
                f"e.g. `{self.prefix} mapzoom USA`."
            )
            return
        map_cfg = self.cfg.get("spatial_mapping", {})
        buffer_deg = float(map_cfg.get("zoom_buffer_deg", 5.0))
        occupation = self.map_store.get_occupation()
        async with message.channel.typing():
            try:
                img_bytes = await map_renderer.render_zoom_map(
                    tag, occupation, buffer_deg=buffer_deg, **self._province_render_args()
                )
            except ValueError as e:
                await message.reply(str(e))
                return
            except Exception as e:
                await message.reply(f"Map render failed: {e}")
                return
        await message.reply(
            file=discord.File(fp=io.BytesIO(img_bytes), filename=f"dnc_map_{tag}.png")
        )

    async def _cmd_mapfaction(self, message: discord.Message) -> None:
        """!DNC mapfaction — render world map colored by faction membership."""
        self.faction_store.reload()
        memberships = self.faction_store.get_memberships()
        if not memberships:
            await message.reply(
                f"No faction memberships are configured. Edit `{self.faction_store.path}` "
                "and add sections like `[NATO]` with `members = USA, GBR, FRA`."
            )
            return

        async with message.channel.typing():
            try:
                img_bytes = await map_renderer.render_faction_map(memberships, year=self.state.current_year)
            except Exception as e:
                await message.reply(f"Faction map render failed: {e}")
                return

        warning_text = ""
        warnings = self.faction_store.warnings
        if warnings:
            warning_text = "\nWarnings: " + "; ".join(warnings[:5])
            if len(warnings) > 5:
                warning_text += f"; +{len(warnings) - 5} more"
        await message.reply(
            f"Faction map rendered from `{self.faction_store.path}`.{warning_text}",
            file=discord.File(fp=io.BytesIO(img_bytes), filename="dnc_faction_map.png"),
        )

    # ── Province command handlers ─────────────────────────────────────────────

    async def _cmd_mapdivide(self, message: discord.Message, arg: str) -> None:
        """!DNC mapdivide — provinces are now always visible from ArcGIS data."""
        tag = arg.strip().upper().split()[0] if arg.strip() else ""
        if tag:
            iso2 = map_renderer.TAG_TO_ISO2.get(tag)
            if not iso2:
                await message.reply(f"`{tag}` has no ArcGIS province mapping. Check `TAG_TO_ISO2`.")
                return
            provinces = self.province_store.list_by_country(tag)
            await message.reply(
                f"`{tag}` has **{len(provinces)}** real ArcGIS provinces (ISO2: `{iso2}`).\n"
                f"Province IDs look like `{provinces[0].id}` … `{provinces[-1].id}`\n"
                "Use `!DNC maplist TAG` to see them all, or `!DNC mapset ID @player` to assign ownership."
                if provinces else f"`{tag}` maps to ISO2 `{iso2}` but no provinces found in ArcGIS data."
            )
        else:
            await message.reply(
                "Provinces are now loaded automatically from real ArcGIS administrative divisions.\n"
                "Use `!DNC maplist TAG` to browse provinces for a country,\n"
                "or `!DNC mapset PROVINCE_ID @player` (e.g. `US-AK`) to assign ownership."
            )

    async def _cmd_mapmerge(self, message: discord.Message, tag: str) -> None:
        """!DNC mapmerge TAG — release all province ownership for TAG."""
        tag = tag.strip().upper()
        if not re.match(r"^[A-Z]{2,4}$", tag):
            await message.reply(f"Usage: `{self.prefix} mapmerge TAG`")
            return

        released = self.province_store.release_all_for_tag(tag)
        if released == 0:
            await message.reply(f"`{tag}` has no owned provinces to release.")
            return

        async with message.channel.typing():
            try:
                img_bytes, _key, _hit = await self._render_world_via_cache()
            except Exception as e:
                await message.reply(f"Released {released} province(s) for `{tag}` (render failed: {e}).")
                return

        self.flog.log_command("mapmerge", str(message.author), message.channel.name, tag)
        await message.reply(
            f"Released **{released}** province(s) for `{tag}` — all now uncontrolled.",
            file=discord.File(fp=io.BytesIO(img_bytes), filename="dnc_map.png"),
        )

    async def _cmd_mapset(self, message: discord.Message, arg: str) -> None:
        """!DNC mapset PROVINCE_ID @player — manually assign province ownership."""
        parts = arg.strip().split(None, 1)
        if len(parts) != 2:
            await message.reply(
                f"Usage: `{self.prefix} mapset PROVINCE_ID @player`\n"
                "Example: `!DNC mapset US-AK @player` (use `!DNC maplist USA` to see IDs)"
            )
            return
        province_id = parts[0].upper()
        mention_part = parts[1].strip()

        m = MENTION_RE.match(mention_part)
        if not m:
            await message.reply("Please @mention the player: `!DNC mapset US-AK @player`")
            return
        user_id = m.group(1)

        member = message.guild.get_member(int(user_id)) if message.guild else None
        display_name = member.display_name if member else f"User#{user_id}"
        player_tag = parse_tag(member.nick or member.name) if member else ""
        if not player_tag:
            # Fall back to first occupied game tag that maps to this province's ISO2
            prov = self.province_store.get_province(province_id)
            if prov:
                for tag in map_renderer._ISO2_TO_TAGS.get(prov.iso_cc, []):
                    if tag in self.map_store.get_occupation():
                        player_tag = tag
                        break
            if not player_tag:
                player_tag = province_id.split("-")[0]

        try:
            self.province_store.assign(province_id, player_tag, user_id, display_name)
        except ValueError as e:
            await message.reply(str(e))
            return

        async with message.channel.typing():
            try:
                img_bytes, _key, _hit = await self._render_world_via_cache()
            except Exception as e:
                await message.reply(f"Assigned `{province_id}` to {display_name} (render failed: {e}).")
                return

        self.flog.log_command("mapset", str(message.author), message.channel.name, arg)
        await message.reply(
            f"Assigned `{province_id}` to **{display_name}**.",
            file=discord.File(fp=io.BytesIO(img_bytes), filename="dnc_map.png"),
        )

    async def _cmd_maprelease(self, message: discord.Message, province_id: str) -> None:
        """!DNC maprelease PROVINCE_ID — release province ownership."""
        province_id = province_id.strip().upper()
        if not province_id:
            await message.reply(f"Usage: `{self.prefix} maprelease PROVINCE_ID`")
            return

        removed = self.province_store.release(province_id)
        if not removed:
            await message.reply(f"`{province_id}` has no current owner.")
            return

        async with message.channel.typing():
            try:
                img_bytes, _key, _hit = await self._render_world_via_cache()
            except Exception as e:
                await message.reply(f"Released `{province_id}` (render failed: {e}).")
                return

        self.flog.log_command("maprelease", str(message.author), message.channel.name, province_id)
        await message.reply(
            f"Released `{province_id}` — now uncontrolled.",
            file=discord.File(fp=io.BytesIO(img_bytes), filename="dnc_map.png"),
        )

    async def _cmd_maplist(self, message: discord.Message, tag: str) -> None:
        """!DNC maplist TAG — list real ArcGIS provinces for TAG and their owners."""
        tag = tag.strip().upper() if tag.strip() else ""

        if not tag:
            # Show countries with owned provinces
            ownership = self.province_store.get_ownership()
            if not ownership:
                await message.reply("No provinces are currently owned. Use `!DNC mapset PROVINCE_ID @player` to assign.")
                return
            by_iso2: dict[str, list] = {}
            for pid, o in ownership.items():
                iso2 = pid.split("-")[0] if "-" in pid else pid
                by_iso2.setdefault(iso2, []).append((pid, o))
            lines = [f"**Owned provinces** ({len(ownership)} total):\n"]
            for iso2, items in sorted(by_iso2.items()):
                lines.append(f"  `{iso2}`: {len(items)} province(s)")
            await message.reply("\n".join(lines))
            return

        provinces = self.province_store.list_by_country(tag)
        if not provinces:
            iso2 = map_renderer.TAG_TO_ISO2.get(tag)
            if not iso2:
                await message.reply(f"`{tag}` has no ArcGIS mapping. Check `TAG_TO_ISO2` in map_renderer.py.")
            else:
                await message.reply(f"`{tag}` (ISO2: `{iso2}`) has no provinces in the ArcGIS dataset.")
            return

        ownership = self.province_store.get_ownership()
        lines = [f"**Provinces of {tag}** ({len(provinces)} total):\n"]
        for p in provinces:
            o = ownership.get(p.id)
            owner_str = f"**{o.display_name}** ({o.player_tag})" if o else "*uncontrolled*"
            lines.append(f"`{p.id}` {p.name} — {owner_str}")

        # Discord message limit: split if too long
        text = "\n".join(lines)
        if len(text) > 1900:
            await message.reply(text[:1900] + f"\n… (+{len(provinces) - text[:1900].count(chr(10))} more)")
        else:
            await message.reply(text)

    async def _cmd_mapparse(self, message: discord.Message, arg: str) -> None:
        """!DNC mapparse <message-link> — LLM reads a post and proposes ownership changes."""
        arg = arg.strip()
        m = MSG_LINK_RE.search(arg)
        if not m:
            await message.reply(
                f"Usage: `{self.prefix} mapparse <discord-message-link>`"
            )
            return

        guild_id, channel_id, msg_id = m.group(1), m.group(2), m.group(3)
        try:
            target_channel = self.get_channel(int(channel_id))
            if target_channel is None:
                await message.reply("Cannot access that channel.")
                return
            target_msg = await target_channel.fetch_message(int(msg_id))
        except discord.NotFound:
            await message.reply("Message not found.")
            return
        except Exception as e:
            await message.reply(f"Failed to fetch message: {e}")
            return

        post_text = target_msg.content or ""
        if not post_text:
            await message.reply("That message has no text content.")
            return

        # Province context is now just the current ownership snapshot — the
        # LLM uses geometry tools to enumerate provinces, so we no longer
        # need to dump the full province list into the prompt.
        ownership = self.province_store.get_ownership()
        province_context = {
            pid: {"owner": o.player_tag, "display_name": o.display_name}
            for pid, o in ownership.items()
        }
        occupation_context = {
            tag: {"display_name": n.display_name, "color": n.color}
            for tag, n in self.map_store.get_occupation().items()
        }

        map_cfg = self.cfg.get("spatial_mapping", {})
        model = map_cfg.get("map_llm_model") or None
        thinking_budget = map_cfg.get("map_llm_thinking_budget")
        max_tool_depth = int(map_cfg.get("map_llm_max_tool_depth", 6) or 6)
        tavily_enabled = (
            self.tavily is not None
            and self.cfg.get("tavily", {}).get("enabled_modes", {}).get("mapparse", False)
        )

        async with message.channel.typing():
            try:
                changes = await map_llm.parse_post_for_changes(
                    post_text, province_context, self.llm, model=model,
                    tavily=self.tavily if tavily_enabled else None,
                    thinking_budget=thinking_budget,
                    max_tool_depth=max_tool_depth,
                    occupation_context=occupation_context,
                )
            except Exception as e:
                await message.reply(f"LLM analysis failed: {e}")
                return

        valid_changes, invalid_changes = map_llm.validate_proposals(changes)
        summary = map_llm.format_change_summary(valid_changes, self.province_store, invalid_changes)
        if len(summary) > 4000:
            summary = summary[:3950] + "\n…*(summary truncated)*"
        confirm_msg = await message.reply(summary)
        changes = valid_changes

        if not changes:
            return

        await confirm_msg.add_reaction("✅")
        await confirm_msg.add_reaction("❌")

        def _check(reaction, user):
            return (
                reaction.message.id == confirm_msg.id
                and str(reaction.emoji) in ("✅", "❌")
                and self._is_admin_user_by_id(user, message.guild)
            )

        try:
            reaction, _ = await self.wait_for("reaction_add", timeout=60.0, check=_check)
        except asyncio.TimeoutError:
            await confirm_msg.reply("Timed out — no changes applied.")
            return

        if str(reaction.emoji) == "❌":
            await confirm_msg.reply("Cancelled — no changes applied.")
            return

        applied, errors = self._apply_map_proposals(
            changes, command="mapparse", actor=str(message.author),
        )

        self.flog.log_command("mapparse", str(message.author), message.channel.name, arg)

        result_lines = [f"Applied **{applied}** change(s)."]
        if errors:
            result_lines.append("Errors: " + "; ".join(errors))

        async with message.channel.typing():
            try:
                img_bytes, _key, _hit = await self._render_world_via_cache()
                await confirm_msg.reply(
                    "\n".join(result_lines),
                    file=discord.File(fp=io.BytesIO(img_bytes), filename="dnc_map.png"),
                )
            except Exception as e:
                await confirm_msg.reply("\n".join(result_lines) + f"\n(Render failed: {e})")

    def _apply_map_proposals(
        self, changes: list[dict], *, command: str, actor: str,
    ) -> tuple[int, list[str]]:
        """Apply a validated list of map-change proposals.

        Returns (applied_count, errors). Each applied operation is recorded
        to logs/map_changes.log.
        """
        applied = 0
        errors: list[str] = []
        for ch in changes:
            action = ch.get("action")
            pid = (ch.get("province_id") or "").upper()
            reason = ch.get("reason", "")
            try:
                if action == "assign_province":
                    player_tag = (ch.get("player_tag") or "").upper()
                    self.province_store.assign(pid, player_tag, "", player_tag)
                    self.flog.log_map_change(
                        command, actor, "assign", pid,
                        player_tag=player_tag, reason=reason,
                    )
                    applied += 1
                elif action == "release_province":
                    self.province_store.release(pid)
                    self.flog.log_map_change(
                        command, actor, "release", pid, reason=reason,
                    )
                    applied += 1
                elif action == "transfer_province":
                    to_tag = (ch.get("to_tag") or "").upper()
                    from_tag = (ch.get("from_tag") or "").upper()
                    self.province_store.assign(pid, to_tag, "", to_tag)
                    self.flog.log_map_change(
                        command, actor, "transfer", pid,
                        from_tag=from_tag, to_tag=to_tag, reason=reason,
                    )
                    applied += 1
            except Exception as e:
                errors.append(f"`{pid}`: {e}")
        return applied, errors

    async def _cmd_mapchange(self, message: discord.Message, arg: str) -> None:
        """!DNC mapchange <instructions> — LLM applies direct map ownership changes."""
        instructions = arg.strip()
        if not instructions:
            await message.reply(
                f"Usage: `{self.prefix} mapchange <instructions>`\n"
                f"Example: `{self.prefix} mapchange Russia invades northern China 450 miles from Xinjiang`"
            )
            return

        map_cfg = self.cfg.get("spatial_mapping", {})
        ownership = self.province_store.get_ownership()
        # Hand the LLM only the current ownership snapshot — it discovers the
        # province universe through geometry tools instead of a flat dump.
        province_context = {
            pid: {"owner": o.player_tag, "display_name": o.display_name}
            for pid, o in ownership.items()
        }
        occupation_context = {
            tag: {"display_name": n.display_name, "color": n.color}
            for tag, n in self.map_store.get_occupation().items()
        }

        model = map_cfg.get("map_llm_model") or None
        thinking_budget = map_cfg.get("map_llm_thinking_budget")
        max_tool_depth = int(map_cfg.get("map_llm_max_tool_depth", 6) or 6)
        tavily_enabled = (
            self.tavily is not None
            and self.cfg.get("tavily", {}).get("enabled_modes", {}).get("mapchange", False)
        )

        async with message.channel.typing():
            try:
                changes = await map_llm.process_map_instructions(
                    instructions, province_context, self.llm, model=model,
                    tavily=self.tavily if tavily_enabled else None,
                    thinking_budget=thinking_budget,
                    max_tool_depth=max_tool_depth,
                    occupation_context=occupation_context,
                )
            except Exception as e:
                await message.reply(f"LLM processing failed: {e}")
                return

        valid_changes, invalid_changes = map_llm.validate_proposals(changes)
        summary = map_llm.format_change_summary(valid_changes, self.province_store, invalid_changes)
        if len(summary) > 4000:
            summary = summary[:3950] + "\n…*(summary truncated)*"
        confirm_msg = await message.reply(summary)
        changes = valid_changes

        if not changes:
            return

        await confirm_msg.add_reaction("✅")
        await confirm_msg.add_reaction("❌")

        def _check(reaction, user):
            return (
                reaction.message.id == confirm_msg.id
                and str(reaction.emoji) in ("✅", "❌")
                and self._is_admin_user_by_id(user, message.guild)
            )

        try:
            reaction, _ = await self.wait_for("reaction_add", timeout=60.0, check=_check)
        except asyncio.TimeoutError:
            await confirm_msg.reply("Timed out — no changes applied.")
            return

        if str(reaction.emoji) == "❌":
            await confirm_msg.reply("Cancelled — no changes applied.")
            return

        applied, errors = self._apply_map_proposals(
            changes, command="mapchange", actor=str(message.author),
        )

        self.flog.log_command("mapchange", str(message.author), message.channel.name, instructions[:120])

        result_lines = [f"Applied **{applied}** change(s)."]
        if errors:
            result_lines.append("Errors: " + "; ".join(errors))

        async with message.channel.typing():
            try:
                img_bytes, _key, _hit = await self._render_world_via_cache()
                await confirm_msg.reply(
                    "\n".join(result_lines),
                    file=discord.File(fp=io.BytesIO(img_bytes), filename="dnc_map.png"),
                )
            except Exception as e:
                await confirm_msg.reply("\n".join(result_lines) + f"\n(Render failed: {e})")

    # ------------------------------------------------------------------
    async def close(self):
        for task in list(self._pending_timers.values()):
            task.cancel()
        self._pending_chains.clear()
        self._pending_timers.clear()
        if self._rollover_task.is_running():
            self._rollover_task.cancel()
        if self._map_render_task.is_running():
            self._map_render_task.cancel()
        await self.llm.close()
        if self.tavily:
            await self.tavily.close()
        await super().close()


# ----------------------------------------------------------------------
class LoreCog(commands.Cog):
    def __init__(self, bot: LoreBot):
        self.bot = bot

    async def cog_check(self, ctx: commands.Context) -> bool:
        if ctx.command.name in ("optout", "optin", "spy", "counterspy", "unspy", "spylist"):
            return True
        if self.bot.optouts.is_opted_out(str(ctx.author.id)):
            await ctx.reply(
                f"You're currently opted out of the lore archive. "
                f"Use `{self.bot.prefix} optin` to re-enable access."
            )
            return False
        return True

    @commands.command(name="chat")
    async def chat(self, ctx: commands.Context, *, body: str = ""):
        await self.bot._cmd_chat(ctx.message, body)

    @commands.command(name="chatuc")
    async def chatuc(self, ctx: commands.Context, *, body: str = ""):
        if not self.bot._is_chatuc_user(ctx.message): return
        self.bot.flog.log_command("chatuc", str(ctx.author), ctx.channel.name, body)
        await self.bot._cmd_chatuc(ctx.message, body)

    @commands.command(name="gm")
    async def gm(self, ctx: commands.Context, *, arg: str = ""):
        await self.bot._cmd_gm(ctx.message, arg)

    @commands.command(name="war")
    async def war(self, ctx: commands.Context, *, arg: str = ""):
        await self.bot._cmd_war(ctx.message, arg)

    @commands.command(name="optout")
    async def optout(self, ctx: commands.Context):
        await self.bot._cmd_optout(ctx.message)

    @commands.command(name="optin")
    async def optin(self, ctx: commands.Context):
        await self.bot._cmd_optin(ctx.message)

    @commands.command(name="year")
    async def year(self, ctx: commands.Context):
        await ctx.reply(f"📅 The current in-game year is **{self.bot.state.current_year}**.")

    @commands.command(name="whoami")
    async def whoami(self, ctx: commands.Context):
        await self.bot._cmd_whoami(ctx.message)

    @commands.command(name="help")
    async def help(self, ctx: commands.Context):
        await ctx.reply(self.bot._help_text(ctx.message))

    @commands.command(name="spy")
    async def spy(self, ctx: commands.Context, *, arg: str = ""):
        await self.bot._cmd_spy(ctx.message, arg)

    @commands.command(name="counterspy")
    async def counterspy(self, ctx: commands.Context):
        await self.bot._cmd_counterspy(ctx.message)

    @commands.command(name="unspy")
    async def unspy(self, ctx: commands.Context, *, arg: str = ""):
        await self.bot._cmd_unspy(ctx.message, arg)

    @commands.command(name="spylist")
    async def spylist(self, ctx: commands.Context):
        await self.bot._cmd_spylist(ctx.message)

    @commands.command(name="ingest")
    async def ingest(self, ctx: commands.Context, *, arg: str = ""):
        if not self.bot._is_admin_channel(ctx.channel) or not self.bot._is_admin_user(ctx.message): return
        self.bot.flog.log_command("ingest", str(ctx.author), ctx.channel.name, arg)
        await self.bot._cmd_ingest(ctx.message, arg)

    @commands.command(name="void")
    async def void(self, ctx: commands.Context, *, arg: str = ""):
        if not self.bot._is_admin_user(ctx.message): return
        self.bot.flog.log_command("void", str(ctx.author), ctx.channel.name, arg)
        await self.bot._cmd_void(ctx.message, arg)

    @commands.command(name="unvoid")
    async def unvoid(self, ctx: commands.Context, *, arg: str = ""):
        if not self.bot._is_admin_user(ctx.message): return
        self.bot.flog.log_command("unvoid", str(ctx.author), ctx.channel.name, arg)
        await self.bot._cmd_unvoid(ctx.message, arg)

    @commands.command(name="yearset")
    async def yearset(self, ctx: commands.Context, *, arg: str = ""):
        if not self.bot._is_admin_channel(ctx.channel) or not self.bot._is_admin_user(ctx.message): return
        self.bot.flog.log_command("yearset", str(ctx.author), ctx.channel.name, arg)
        await self.bot._cmd_yearset(ctx.message, arg)

    @commands.command(name="yearroll")
    async def yearroll(self, ctx: commands.Context):
        if not self.bot._is_admin_channel(ctx.channel) or not self.bot._is_admin_user(ctx.message): return
        self.bot.flog.log_command("yearroll", str(ctx.author), ctx.channel.name, "")
        await self.bot._cmd_yearroll(ctx.message)

    @commands.command(name="purge")
    async def purge(self, ctx: commands.Context, *, arg: str = ""):
        if not self.bot._is_admin_channel(ctx.channel) or not self.bot._is_admin_user(ctx.message): return
        self.bot.flog.log_command("purge", str(ctx.author), ctx.channel.name, arg)
        await self.bot._cmd_purge(ctx.message, arg)

    @commands.command(name="export")
    async def export(self, ctx: commands.Context):
        if not self.bot._is_admin_channel(ctx.channel) or not self.bot._is_admin_user(ctx.message): return
        self.bot.flog.log_command("export", str(ctx.author), ctx.channel.name, "")
        await self.bot._cmd_export(ctx.message)

    @commands.command(name="stats")
    async def stats(self, ctx: commands.Context, *, arg: str = ""):
        if not self.bot._is_admin_channel(ctx.channel) or not self.bot._is_admin_user(ctx.message): return
        self.bot.flog.log_command("stats", str(ctx.author), ctx.channel.name, arg)
        await self.bot._cmd_stats(ctx.message, arg)

    @commands.command(name="channels")
    async def channels(self, ctx: commands.Context):
        if not self.bot._is_admin_channel(ctx.channel) or not self.bot._is_admin_user(ctx.message): return
        self.bot.flog.log_command("channels", str(ctx.author), ctx.channel.name, "")
        await self.bot._cmd_channels(ctx.message)

    @commands.command(name="reloadprompts")
    async def reloadprompts(self, ctx: commands.Context):
        if not self.bot._is_admin_channel(ctx.channel) or not self.bot._is_admin_user(ctx.message): return
        self.bot.flog.log_command("reloadprompts", str(ctx.author), ctx.channel.name, "")
        await self.bot._cmd_reloadprompts(ctx.message)

    @commands.command(name="chatban")
    async def chatban(self, ctx: commands.Context, *, arg: str = ""):
        if not self.bot._is_admin_channel(ctx.channel) or not self.bot._is_admin_user(ctx.message): return
        self.bot.flog.log_command("chatban", str(ctx.author), ctx.channel.name, arg)
        await self.bot._cmd_chatban(ctx.message, arg)

    @commands.command(name="chatunban")
    async def chatunban(self, ctx: commands.Context, *, arg: str = ""):
        if not self.bot._is_admin_channel(ctx.channel) or not self.bot._is_admin_user(ctx.message): return
        self.bot.flog.log_command("chatunban", str(ctx.author), ctx.channel.name, arg)
        await self.bot._cmd_chatunban(ctx.message, arg)

    @commands.command(name="map")
    async def map(self, ctx: commands.Context):
        await self.bot._cmd_map(ctx.message)

    @commands.command(name="mapzoom")
    async def mapzoom(self, ctx: commands.Context, *, tag: str = ""):
        await self.bot._cmd_mapzoom(ctx.message, tag)

    @commands.command(name="mapfaction")
    async def mapfaction(self, ctx: commands.Context):
        await self.bot._cmd_mapfaction(ctx.message)

    @commands.command(name="mapdivide")
    async def mapdivide(self, ctx: commands.Context, *, arg: str = ""):
        if not self.bot._is_admin_channel(ctx.channel) or not self.bot._is_admin_user(ctx.message): return
        await self.bot._cmd_mapdivide(ctx.message, arg)

    @commands.command(name="mapmerge")
    async def mapmerge(self, ctx: commands.Context, *, tag: str = ""):
        if not self.bot._is_admin_channel(ctx.channel) or not self.bot._is_admin_user(ctx.message): return
        await self.bot._cmd_mapmerge(ctx.message, tag)

    @commands.command(name="mapset")
    async def mapset(self, ctx: commands.Context, *, arg: str = ""):
        if not self.bot._is_admin_channel(ctx.channel) or not self.bot._is_admin_user(ctx.message): return
        await self.bot._cmd_mapset(ctx.message, arg)

    @commands.command(name="maprelease")
    async def maprelease(self, ctx: commands.Context, *, province_id: str = ""):
        if not self.bot._is_admin_channel(ctx.channel) or not self.bot._is_admin_user(ctx.message): return
        await self.bot._cmd_maprelease(ctx.message, province_id)

    @commands.command(name="maplist")
    async def maplist(self, ctx: commands.Context, *, tag: str = ""):
        await self.bot._cmd_maplist(ctx.message, tag)

    @commands.command(name="mapparse")
    async def mapparse(self, ctx: commands.Context, *, arg: str = ""):
        if not self.bot._is_admin_channel(ctx.channel) or not self.bot._is_admin_user(ctx.message): return
        await self.bot._cmd_mapparse(ctx.message, arg)

    @commands.command(name="mapchange")
    async def mapchange(self, ctx: commands.Context, *, arg: str = ""):
        if not self.bot._is_admin_channel(ctx.channel) or not self.bot._is_admin_user(ctx.message): return
        await self.bot._cmd_mapchange(ctx.message, arg)


def main():
    load_dotenv()
    if "DISCORD_TOKEN" not in os.environ:
        sys.exit("DISCORD_TOKEN not set")
    if "OPENROUTER_API_KEY" not in os.environ:
        sys.exit("OPENROUTER_API_KEY not set")

    config = load_config()
    bot = LoreBot(config)
    bot.run(os.environ["DISCORD_TOKEN"], log_handler=None)


if __name__ == "__main__":
    main()
